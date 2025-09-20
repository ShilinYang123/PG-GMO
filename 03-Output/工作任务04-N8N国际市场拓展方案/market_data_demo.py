#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
N8N市场调查数据获取演示脚本
模拟从API获取市场数据的过程
"""

import json
import random
from datetime import datetime
import time
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("警告: python-docx未安装，将使用JSON格式保存报告")
    Document = None

def generate_market_data():
    """生成模拟的市场调查数据"""
    
    # 模拟产品数据
    products = [
        "智能空调", "变频洗衣机", "节能冰箱", "智能热水器", 
        "空气净化器", "智能电视", "微波炉", "电磁炉"
    ]
    
    # 模拟竞争对手
    competitors = [
        ["美的", "格力", "海尔"],
        ["小米", "华为", "TCL"],
        ["松下", "三星", "LG"]
    ]
    
    market_data = []
    
    for i, product in enumerate(products[:5]):  # 获取5个产品的数据
        data = {
            "product_name": product,
            "price": round(random.uniform(1000, 8000), 2),
            "market_share": round(random.uniform(5, 25), 1),
            "competitors": random.choice(competitors),
            "trend": random.choice(["上升", "下降", "稳定"]),
            "region": random.choice(["华东", "华南", "华北", "西南"]),
            "survey_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        market_data.append(data)
        
        # 模拟API请求延迟
        time.sleep(0.5)
        print(f"✓ 已获取 {product} 的市场数据")
    
    return market_data

def validate_data(data):
    """验证数据完整性"""
    print("\n🔍 开始数据验证...")
    
    valid_data = []
    for item in data:
        if all(key in item for key in ["product_name", "price", "market_share"]):
            if item["price"] > 0 and item["market_share"] > 0:
                valid_data.append(item)
                print(f"✓ {item['product_name']} 数据验证通过")
            else:
                print(f"✗ {item['product_name']} 数据异常，已过滤")
        else:
            print(f"✗ 数据字段不完整，已过滤")
    
    return valid_data

def process_data(data):
    """处理和分析数据"""
    print("\n⚙️ 开始数据处理...")
    
    processed_data = []
    total_market_value = 0
    
    for item in data:
        # 计算市场价值
        market_value = item["price"] * item["market_share"] * 100
        total_market_value += market_value
        
        processed_item = {
            "产品名称": item["product_name"],
            "市场价格": f"¥{item['price']:.2f}",
            "市场份额": f"{item['market_share']}%",
            "主要竞争对手": ", ".join(item["competitors"]),
            "价格趋势": item["trend"],
            "主要区域": item["region"],
            "市场价值": f"¥{market_value:,.2f}",
            "调查时间": item["survey_date"]
        }
        processed_data.append(processed_item)
        print(f"✓ 已处理 {item['product_name']} 数据")
    
    # 添加汇总信息
    summary = {
        "总市场价值": f"¥{total_market_value:,.2f}",
        "调查产品数量": len(processed_data),
        "平均市场份额": f"{sum(float(item['市场份额'].rstrip('%')) for item in processed_data) / len(processed_data):.1f}%"
    }
    
    return processed_data, summary

def save_to_word(data, summary, filename):
    """保存数据到Word文件"""
    if Document is None:
        # 如果没有安装python-docx，回退到JSON格式
        save_to_json(data, summary, filename.replace('.docx', '.json'))
        return
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('市场调查数据报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加生成时间
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    doc.add_paragraph('')
    
    # 添加数据汇总
    doc.add_heading('📊 数据汇总', level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '数值'
    
    for key, value in summary.items():
        row_cells = summary_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
    
    doc.add_paragraph('')
    
    # 添加详细数据
    doc.add_heading('📋 详细数据', level=1)
    
    for i, item in enumerate(data, 1):
        doc.add_heading(f'{i}. {item["产品名称"]}', level=2)
        
        # 创建产品信息表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '属性'
        hdr_cells[1].text = '值'
        
        for key, value in item.items():
            if key != '产品名称':  # 产品名称已经作为标题
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
        
        doc.add_paragraph('')
    
    # 添加页脚
    doc.add_paragraph('')
    footer = doc.add_paragraph('本报告由N8N自动化工作流生成')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(filename)
    print(f"\n💾 Word报告已保存到: {filename}")

def save_to_json(data, summary, filename):
    """保存数据到JSON文件（备用方法）"""
    report = {
        "报告标题": "市场调查数据报告",
        "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "数据汇总": summary,
        "详细数据": data
    }
    
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    
    print(f"\n💾 数据已保存到: {filename}")

def main():
    """主函数 - 模拟N8N工作流执行过程"""
    print("🚀 N8N市场调查工作流演示")
    print("=" * 50)
    
    # 步骤1: 获取市场数据
    print("\n📊 步骤1: 获取市场数据")
    raw_data = generate_market_data()
    
    # 步骤2: 数据验证
    print("\n✅ 步骤2: 数据验证")
    valid_data = validate_data(raw_data)
    
    # 步骤3: 数据处理
    print("\n⚙️ 步骤3: 数据处理")
    processed_data, summary = process_data(valid_data)
    
    # 步骤4: 保存报告
    print("\n📋 步骤4: 生成报告")
    report_filename = f"S:\\PG-GMO\\02-Output\\N8N国际市场拓展方案\\市场调查报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    save_to_word(processed_data, summary, report_filename)
    
    # 步骤5: 显示结果
    print("\n📈 步骤5: 调查结果")
    print("-" * 30)
    for key, value in summary.items():
        print(f"{key}: {value}")
    
    print("\n📧 步骤6: 发送通知")
    print("✓ 市场调查完成通知已发送")
    print(f"✓ 报告文件: {report_filename}")
    
    print("\n🎉 工作流执行完成！")
    return report_filename

if __name__ == "__main__":
    report_file = main()