#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
甯傚満璋冩煡鎶ュ憡Word鏍煎紡鐢熸垚鍣?灏哅arkdown鏍煎紡鐨勫競鍦鸿皟鏌ユ姤鍛婅浆鎹负Word鏍煎紡
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_hyperlink(paragraph, text, url):
    """鍦ㄦ钀戒腑娣诲姞瓒呴摼鎺?""
    # 鑾峰彇娈佃惤鐨勭埗鏂囨。
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # 鍒涘缓瓒呴摼鎺ュ厓绱?    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # 鍒涘缓杩愯鍏冪礌
    new_run = OxmlElement('w:r')
    
    # 璁剧疆瓒呴摼鎺ユ牱寮?    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # 娣诲姞鏂囨湰
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_word_report():
    """鍒涘缓Word鏍煎紡鐨勫競鍦鸿皟鏌ユ姤鍛?""
    
    # 璇诲彇Markdown鎶ュ憡鍐呭
    md_file_path = "S:\\PG-GMO\\\02-Output\\\甯傚満璋冩煡\\鍝侀珮鐢靛櫒\\宸ヤ綔浠诲姟-鍝侀珮2508.md"
    
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"閿欒锛氭壘涓嶅埌鏂囦欢 {md_file_path}")
        return
    
    # 鍒涘缓Word鏂囨。
    doc = Document()
    
    # 璁剧疆鏂囨。鏍囬
    title = doc.add_heading('姹熼棬甯傚搧楂樼數鍣ㄥ疄涓氭湁闄愬叕鍙稿皬瀹剁數甯傚満娣卞害璋冩煡鎶ュ憡', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 瑙ｆ瀽Markdown鍐呭骞惰浆鎹负Word鏍煎紡
    lines = content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # 澶勭悊鏍囬
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # 澶勭悊鍒楄〃椤?        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # 澶勭悊鏁板瓧鍒楄〃
        elif line.startswith(tuple(f'{i}. ' for i in range(1, 10))):
            p = doc.add_paragraph(line[3:], style='List Number')
        
        # 澶勭悊绮椾綋鏂囨湰
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    p.add_run(part).bold = True
        
        # 澶勭悊鍖呭惈閾炬帴鐨勬枃鏈?        elif 'http' in line and ('鏉ユ簮' in line or '鏁版嵁' in line or '鍙傝€? in line):
            p = doc.add_paragraph()
            
            # 绠€鍗曠殑URL鎻愬彇鍜岄摼鎺ュ垱寤?            import re
            url_pattern = r'https?://[^\s]+'
            urls = re.findall(url_pattern, line)
            
            if urls:
                # 鍒嗗壊鏂囨湰鍜孶RL
                text_parts = re.split(url_pattern, line)
                url_index = 0
                
                for i, text_part in enumerate(text_parts):
                    if text_part:
                        p.add_run(text_part)
                    if url_index < len(urls) and i < len(text_parts) - 1:
                        # 娣诲姞瓒呴摼鎺?                        add_hyperlink(p, urls[url_index], urls[url_index])
                        url_index += 1
            else:
                p.add_run(line)
        
        # 澶勭悊鏅€氭钀?        elif line and not line.startswith('#'):
            # 璺宠繃Markdown鍏冩暟鎹?            if line.startswith('---') or line.startswith('璋冩煡鏃ユ湡:') or line.startswith('璋冩煡浜哄憳:'):
                continue
            
            p = doc.add_paragraph(line)
    
    # 娣诲姞椤佃剼
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "姹熼棬甯傚搧楂樼數鍣ㄥ疄涓氭湁闄愬叕鍙?- 甯傚満璋冩煡鎶ュ憡"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 淇濆瓨Word鏂囨。
    output_path = "S:\\PG-GMO\\\02-Output\\\甯傚満璋冩煡\\鍝侀珮鐢靛櫒\\鍝侀珮鐢靛櫒灏忓鐢靛競鍦鸿皟鏌ユ姤鍛?docx"
    doc.save(output_path)
    
    print(f"鉁?Word鏍煎紡鎶ュ憡宸茬敓鎴愶細{output_path}")
    print("馃搵 鎶ュ憡鐗圭偣锛?)
    print("   - 鍖呭惈瀹屾暣鐨勫競鍦鸿皟鏌ュ唴瀹?)
    print("   - 鏁版嵁鏉ユ簮閾炬帴鍙偣鍑昏烦杞?)
    print("   - 涓撲笟鐨刉ord鏂囨。鏍煎紡")
    print("   - 閫傚悎鎵撳嵃鍜屽垎浜?)
    
    return output_path

if __name__ == "__main__":
    create_word_report()
