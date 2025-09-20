#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
鏀硅繘鐨刉ord鏂囨。鏍煎紡鎺掔増鍣?浣滆€咃細闆ㄤ繆
鍔熻兘锛氫紭鍖朠G绠＄悊鎵嬪唽鐨勬牸寮忔帓鐗堬紝鎻愬崌鏂囨。涓撲笟鎬у拰鍙鎬?"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import re

def setup_document_styles(doc):
    """璁剧疆鏂囨。鏍峰紡"""
    # 鑾峰彇鏍峰紡闆嗗悎
    styles = doc.styles
    
    # 璁剧疆姝ｆ枃鏍峰紡
    try:
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '瀹嬩綋'
        normal_font.size = Pt(12)
        
        # 璁剧疆娈佃惤鏍煎紡
        normal_para = normal_style.paragraph_format
        normal_para.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_para.line_spacing = 1.5  # 1.5鍊嶈璺?        normal_para.space_after = Pt(6)  # 娈靛悗闂磋窛
        normal_para.first_line_indent = Pt(24)  # 棣栬缂╄繘2瀛楃
    except:
        pass
    
    # 璁剧疆鏍囬鏍峰紡
    try:
        # 涓€绾ф爣棰?        heading1_style = styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = '寰蒋闆呴粦'
        heading1_font.size = Pt(18)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(0, 0, 0)
        
        heading1_para = heading1_style.paragraph_format
        heading1_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading1_para.space_before = Pt(18)
        heading1_para.space_after = Pt(12)
        heading1_para.keep_with_next = True
        
        # 浜岀骇鏍囬
        heading2_style = styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = '寰蒋闆呴粦'
        heading2_font.size = Pt(16)
        heading2_font.bold = True
        heading2_font.color.rgb = RGBColor(0, 0, 0)
        
        heading2_para = heading2_style.paragraph_format
        heading2_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading2_para.space_before = Pt(12)
        heading2_para.space_after = Pt(6)
        heading2_para.keep_with_next = True
        
        # 涓夌骇鏍囬
        heading3_style = styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = '寰蒋闆呴粦'
        heading3_font.size = Pt(14)
        heading3_font.bold = True
        heading3_font.color.rgb = RGBColor(0, 0, 0)
        
        heading3_para = heading3_style.paragraph_format
        heading3_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading3_para.space_before = Pt(12)
        heading3_para.space_after = Pt(6)
        heading3_para.first_line_indent = Pt(0)
        
    except Exception as e:
        print(f"璁剧疆鏍囬鏍峰紡鏃跺嚭閿? {e}")

def create_professional_cover(doc, title="PG绠＄悊鎵嬪唽", subtitle="閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉姹囩紪"):
    """鍒涘缓涓撲笟灏侀潰"""
    # 璁剧疆灏侀潰椤佃竟璺?    section = doc.sections[0]
    section.top_margin = Inches(2)
    section.bottom_margin = Inches(2)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    
    # 涓绘爣棰?    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = '寰蒋闆呴粦'
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # 娣辫摑鑹?    
    # 娣诲姞瑁呴グ绾?    line_para = doc.add_paragraph()
    line_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line_run = line_para.add_run('鈹? * 20)
    line_run.font.name = '寰蒋闆呴粦'
    line_run.font.size = Pt(14)
    line_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 娣诲姞绌鸿
    for _ in range(8):
        doc.add_paragraph()
    
    # 鍓爣棰?    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.name = '寰蒋闆呴粦'
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)  # 鐏拌壊
    
    # 娣诲姞绌鸿
    for _ in range(12):
        doc.add_paragraph()
    
    # 鏃ユ湡鍜岀増鏈俊鎭?    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run(f'{datetime.now().strftime("%Y骞?m鏈?)} 鐗堟湰')
    date_run.font.name = '寰蒋闆呴粦'
    date_run.font.size = Pt(16)
    date_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # 娣诲姞鍒嗛〉绗?    doc.add_page_break()

def create_enhanced_toc(doc, toc_entries):
    """鍒涘缓澧炲己鐗堢洰褰?""
    # 鐩綍鏍囬
    toc_title = doc.add_heading('鐩綍', level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞鐩綍璇存槑
    desc_para = doc.add_paragraph()
    desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    desc_run = desc_para.add_run('Contents')
    desc_run.font.name = 'Times New Roman'
    desc_run.font.size = Pt(12)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # 绌鸿
    
    # 娣诲姞鐩綍鏉＄洰
    for entry in toc_entries:
        level = entry['level']
        title = entry['title']
        page = entry['page']
        
        toc_para = doc.add_paragraph()
        
        # 鏍规嵁绾у埆璁剧疆鏍峰紡
        if level == 1:
            indent = 0
            font_size = Pt(14)
            is_bold = True
            color = RGBColor(0, 51, 102)
        elif level == 2:
            indent = 0.5
            font_size = Pt(12)
            is_bold = False
            color = RGBColor(0, 0, 0)
        else:
            indent = 1.0
            font_size = Pt(11)
            is_bold = False
            color = RGBColor(102, 102, 102)
        
        # 璁剧疆娈佃惤鏍煎紡
        toc_para.paragraph_format.left_indent = Inches(indent)
        toc_para.paragraph_format.space_after = Pt(3)
        
        # 娣诲姞鏍囬
        title_run = toc_para.add_run(title)
        title_run.font.name = '寰蒋闆呴粦'
        title_run.font.size = font_size
        title_run.font.bold = is_bold
        title_run.font.color.rgb = color
        
        # 娣诲姞鐐圭嚎
        dots_length = max(5, 50 - len(title) - len(str(page)))
        dots_run = toc_para.add_run('路' * dots_length)
        dots_run.font.name = '寰蒋闆呴粦'
        dots_run.font.size = Pt(10)
        dots_run.font.color.rgb = RGBColor(153, 153, 153)
        
        # 娣诲姞椤电爜
        page_run = toc_para.add_run(str(page))
        page_run.font.name = 'Times New Roman'
        page_run.font.size = font_size
        page_run.font.bold = is_bold
        page_run.font.color.rgb = color
    
    # 娣诲姞鍒嗛〉绗?    doc.add_page_break()

def setup_page_layout(doc):
    """璁剧疆椤甸潰甯冨眬"""
    for section in doc.sections:
        # 璁剧疆椤佃竟璺?        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
        
        # 璁剧疆椤电湁椤佃剼
        header = section.header
        footer = section.footer
        
        # 娓呯┖榛樿鍐呭
        for para in header.paragraphs:
            para.clear()
        for para in footer.paragraphs:
            para.clear()
        
        # 璁剧疆椤佃剼椤电爜
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 娣诲姞椤电爜
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # 璁剧疆椤电爜瀛椾綋
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

def format_paragraph_content(para, new_para):
    """鏍煎紡鍖栨钀藉唴瀹?""
    if not para.text.strip():
        return
    
    # 妫€鏌ユ槸鍚︽槸鍒楄〃椤?    text = para.text.strip()
    if text.startswith(('鈥?, '路', '-', '*')) or re.match(r'^\d+[.)銆乚', text):
        # 鍒楄〃椤规牸寮?        new_para.paragraph_format.left_indent = Inches(0.5)
        new_para.paragraph_format.first_line_indent = Inches(-0.25)
        new_para.paragraph_format.space_after = Pt(3)
    else:
        # 鏅€氭钀芥牸寮?        new_para.paragraph_format.first_line_indent = Pt(24)  # 棣栬缂╄繘
        new_para.paragraph_format.space_after = Pt(6)
    
    # 澶嶅埗鏂囨湰鍜屾牸寮?    for run in para.runs:
        new_run = new_para.add_run(run.text)
        if run.font.name:
            new_run.font.name = run.font.name
        else:
            new_run.font.name = '瀹嬩綋'
        
        if run.font.size:
            new_run.font.size = run.font.size
        else:
            new_run.font.size = Pt(12)
        
        new_run.font.bold = run.font.bold
        new_run.font.italic = run.font.italic
        new_run.font.underline = run.font.underline

def improve_document_formatting(input_file, output_file):
    """鏀硅繘鏂囨。鏍煎紡"""
    try:
        # 鎵撳紑鍘熸枃妗?        doc = Document(input_file)
        
        # 鍒涘缓鏂版枃妗?        new_doc = Document()
        
        # 璁剧疆鏂囨。鏍峰紡
        setup_document_styles(new_doc)
        
        # 璁剧疆椤甸潰甯冨眬
        setup_page_layout(new_doc)
        
        print("姝ｅ湪鏀硅繘鏂囨。鏍煎紡...")
        
        # 鏀堕泦鐩綍淇℃伅
        toc_entries = []
        current_page = 3
        
        # 鍒嗘瀽鍘熸枃妗ｇ粨鏋?        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                toc_entries.append({
                    'level': level,
                    'title': para.text,
                    'page': current_page
                })
                current_page += 1
        
        # 鍒涘缓灏侀潰
        create_professional_cover(new_doc)
        
        # 鍒涘缓鐩綍
        create_enhanced_toc(new_doc, toc_entries)
        
        # 澶嶅埗骞舵牸寮忓寲鍐呭
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    # 鏍囬
                    level = int(para.style.name.split()[-1])
                    new_heading = new_doc.add_heading(para.text, level=level)
                else:
                    # 鏅€氭钀?                    new_para = new_doc.add_paragraph()
                    format_paragraph_content(para, new_para)
        
        # 澶嶅埗琛ㄦ牸
        for table in doc.tables:
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = 'Table Grid'
            
            # 璁剧疆琛ㄦ牸鏍峰紡
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    new_cell.text = cell.text
                    
                    # 璁剧疆鍗曞厓鏍煎瓧浣?                    for para in new_cell.paragraphs:
                        for run in para.runs:
                            run.font.name = '瀹嬩綋'
                            run.font.size = Pt(10)
        
        # 淇濆瓨鏂囨。
        new_doc.save(output_file)
        print(f"鉁?鏍煎紡鏀硅繘瀹屾垚: {output_file}")
        return True
        
    except Exception as e:
        print(f"鉂?鏍煎紡鏀硅繘澶辫触: {e}")
        return False

def main():
    """涓诲嚱鏁?""
    print("=== Word鏂囨。鏍煎紡鎺掔増鏀硅繘鍣?===")
    print("浣滆€咃細闆ㄤ繆")
    print("寮€濮嬫敼杩涙枃妗ｆ牸寮?..\n")
    
    # 杈撳叆鍜岃緭鍑烘枃浠?    input_file = "S:/PG-GMO/02-Output/PG绠＄悊鎵嬪唽.docx"
    output_file = "S:/PG-GMO/02-Output/PG绠＄悊鎵嬪唽_鏀硅繘鐗?docx"
    
    if not os.path.exists(input_file):
        print(f"閿欒锛氳緭鍏ユ枃浠朵笉瀛樺湪 {input_file}")
        return
    
    # 寮€濮嬫敼杩?    success = improve_document_formatting(input_file, output_file)
    
    if success:
        print("\n馃帀 鏂囨。鏍煎紡鏀硅繘瀹屾垚锛?)
        print(f"鏀硅繘鍚庢枃浠? {output_file}")
        print("\n鏀硅繘鍐呭:")
        print("- 涓撲笟灏侀潰璁捐")
        print("- 澧炲己鐗堢洰褰曟牸寮?)
        print("- 缁熶竴瀛椾綋鍜岄棿璺?)
        print("- 浼樺寲鏍囬鏍峰紡")
        print("- 鏀硅繘娈佃惤鏍煎紡")
        print("- 缇庡寲琛ㄦ牸鏍峰紡")
        print("- 瑙勮寖椤电爜鏄剧ず")
    else:
        print("\n鉂?鏂囨。鏍煎紡鏀硅繘澶辫触")

if __name__ == "__main__":
    main()
