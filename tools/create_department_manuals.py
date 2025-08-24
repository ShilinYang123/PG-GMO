#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
部门管理手册生成器
作者：雨俊
功能：为每个部门单独生成一个Word文档，格式为'PG管理手册-XX部.docx'
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import re

def create_cover_page(doc, department_name):
    """创建部门封面页"""
    # 添加封面标题
    title = doc.add_heading('', level=0)
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.text = f'PG管理手册'
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 添加部门名称
    dept_title = doc.add_heading('', level=0)
    dept_run = dept_title.runs[0] if dept_title.runs else dept_title.add_run()
    dept_run.text = f'{department_name}'
    dept_run.font.name = '微软雅黑'
    dept_run.font.size = Pt(28)
    dept_run.font.bold = True
    dept_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加空行
    for _ in range(6):
        doc.add_paragraph()
    
    # 添加副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('部门基础建设工作成果')
    subtitle_run.font.name = '微软雅黑'
    subtitle_run.font.size = Pt(20)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加空行
    for _ in range(8):
        doc.add_paragraph()
    
    # 添加日期
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'{datetime.now().strftime("%Y年%m月")}')
    date_run.font.name = '微软雅黑'
    date_run.font.size = Pt(16)
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加分页符
    doc.add_page_break()

def create_department_toc(doc, file_list):
    """创建部门目录页"""
    # 目录标题
    toc_title = doc.add_heading('目录', level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # 添加目录条目
    page_num = 3  # 从第3页开始
    for i, file_name in enumerate(file_list, 1):
        toc_para = doc.add_paragraph()
        
        # 设置段落格式
        toc_para.paragraph_format.left_indent = Inches(0.5)
        
        # 添加序号和标题
        title_text = f"{i}. {file_name}"
        title_run = toc_para.add_run(title_text)
        title_run.font.name = '宋体'
        title_run.font.size = Pt(12)
        
        # 添加点线
        dots_count = max(1, 50 - len(title_text) - len(str(page_num)))
        dots_run = toc_para.add_run('.' * dots_count)
        dots_run.font.name = '宋体'
        dots_run.font.size = Pt(12)
        
        # 添加页码
        page_run = toc_para.add_run(str(page_num))
        page_run.font.name = '宋体'
        page_run.font.size = Pt(12)
        
        page_num += 2  # 每个文档估算2页
    
    # 添加分页符
    doc.add_page_break()

def setup_page_numbering(doc):
    """设置页码"""
    try:
        # 获取所有节
        sections = doc.sections
        
        for section in sections:
            # 设置页脚
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加页码字段
            run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)
            
    except Exception as e:
        print(f"设置页码时出错: {e}")

def create_department_manual(dept_dir, output_dir):
    """为单个部门创建管理手册"""
    dept_name = dept_dir.name
    print(f"\n正在处理部门: {dept_name}")
    
    # 获取部门下的所有docx文件
    docx_files = list(dept_dir.rglob('*.docx'))
    if not docx_files:
        print(f"  警告: {dept_name} 没有找到docx文件")
        return False
    
    # 按文件名排序
    docx_files.sort(key=lambda x: x.name)
    
    # 创建主文档
    main_doc = Document()
    
    # 设置文档样式
    try:
        # 设置页面边距
        sections = main_doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    except Exception as e:
        print(f"  设置页面边距时出错: {e}")
    
    # 创建封面
    create_cover_page(main_doc, dept_name)
    
    # 创建目录
    file_names = [f.stem for f in docx_files]
    create_department_toc(main_doc, file_names)
    
    # 合并部门文档
    for docx_file in docx_files:
        try:
            # 读取文档
            sub_doc = Document(str(docx_file))
            
            # 获取文件名作为章节标题
            file_title = docx_file.stem
            
            # 添加章节标题
            chapter_title = main_doc.add_heading(file_title, level=1)
            chapter_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 复制文档内容
            for para in sub_doc.paragraphs:
                if para.text.strip():  # 跳过空段落
                    new_para = main_doc.add_paragraph()
                    
                    # 复制段落格式和内容
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        # 复制字体格式
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                    
                    # 复制段落样式
                    if para.style.name.startswith('Heading'):
                        new_para.style = para.style
                    new_para.alignment = para.alignment
            
            # 添加表格
            for table in sub_doc.tables:
                new_table = main_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = 'Table Grid'
                
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
            
            # 文档之间添加分页
            main_doc.add_page_break()
            
            print(f"  已合并: {docx_file.name}")
            
        except Exception as e:
            print(f"  合并文件失败 {docx_file}: {e}")
            continue
    
    # 设置页码
    setup_page_numbering(main_doc)
    
    # 保存文档
    try:
        output_file = output_dir / f"PG管理手册-{dept_name}.docx"
        os.makedirs(output_dir, exist_ok=True)
        main_doc.save(str(output_file))
        print(f"  ✅ 生成成功: {output_file}")
        return True
    except Exception as e:
        print(f"  ❌ 保存失败: {e}")
        return False

def main():
    """主函数"""
    print("=== 部门管理手册生成器 ===")
    print("作者：雨俊")
    print("开始为每个部门生成独立的管理手册...\n")
    
    # 源目录和输出目录
    source_dir = Path("S:/PG-GMO/Output/部门基础建设工作成果_docx")
    output_dir = Path("S:/PG-GMO/Output")
    
    if not source_dir.exists():
        print(f"错误：源目录不存在 {source_dir}")
        return
    
    # 统计信息
    success_count = 0
    total_count = 0
    
    # 遍历所有部门目录
    for dept_dir in source_dir.iterdir():
        if dept_dir.is_dir():
            total_count += 1
            if create_department_manual(dept_dir, output_dir):
                success_count += 1
    
    # 输出统计结果
    print(f"\n=== 生成完成统计 ===")
    print(f"总部门数: {total_count}")
    print(f"成功生成: {success_count}")
    print(f"失败数量: {total_count - success_count}")
    
    if success_count == total_count:
        print("\n🎉 所有部门管理手册生成成功！")
    else:
        print(f"\n⚠️ 有 {total_count - success_count} 个部门生成失败")
    
    print(f"\n生成的文件保存在: {output_dir}")
    print("文件命名格式: PG管理手册-XX部.docx")

if __name__ == "__main__":
    main()