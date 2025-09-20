import os
import docx
from docx import Document
import PyPDF2
import pandas as pd

def read_word_file(file_path):
    """读取Word文档内容"""
    try:
        doc = Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    content.append(" | ".join(row_data))
        
        return "\n".join(content)
    except Exception as e:
        return f"读取Word文件出错: {str(e)}"

def read_pdf_file(file_path):
    """读取PDF文档内容"""
    try:
        content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    content.append(f"=== 第{page_num + 1}页 ===")
                    content.append(text.strip())
        
        return "\n".join(content)
    except Exception as e:
        return f"读取PDF文件出错: {str(e)}"

# 采购部文件路径
office_files = {
    'Word文档': [
        r'S:\PG-GMO\office\采购部\开发跟单采购员绩效考核方案.docx',
        r'S:\PG-GMO\office\采购部\采购控制程序 HQ-CG-001 A0版.docx',
        r'S:\PG-GMO\office\采购部\采购部部门架构及工作职责 HQ-CG-002 A0版.docx'
    ],
    'PDF文档': [
        r'S:\PG-GMO\office\采购部\开发跟单采购员绩效考核方案.pdf',
        r'S:\PG-GMO\office\采购部\采购控制程序 HQ-CG-001 A0版.pdf',
        r'S:\PG-GMO\office\采购部\采购部部门架构及工作职责 HQ-CG-002 A0版.pdf'
    ]
}

print("开始读取采购部所有Office文件...")
print("=" * 80)

# 读取Word文档
print("\n=== Word文档 ===")
for file_path in office_files['Word文档']:
    if os.path.exists(file_path):
        filename = os.path.basename(file_path)
        print(f"\n--- 文件: {filename} ---")
        content = read_word_file(file_path)
        print(content[:2000])  # 显示前2000个字符
        if len(content) > 2000:
            print(f"\n... 还有 {len(content) - 2000} 个字符 ...")
        print("\n" + "-" * 60)
    else:
        print(f"文件不存在: {file_path}")

# 读取PDF文档
print("\n=== PDF文档 ===")
for file_path in office_files['PDF文档']:
    if os.path.exists(file_path):
        filename = os.path.basename(file_path)
        print(f"\n--- 文件: {filename} ---")
        content = read_pdf_file(file_path)
        print(content[:2000])  # 显示前2000个字符
        if len(content) > 2000:
            print(f"\n... 还有 {len(content) - 2000} 个字符 ...")
        print("\n" + "-" * 60)
    else:
        print(f"文件不存在: {file_path}")

print("\n" + "=" * 80)
print("Office文件读取完成")