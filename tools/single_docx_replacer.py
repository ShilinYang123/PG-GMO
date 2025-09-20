#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
单个DOCX文档内容替换工具
用于处理单个docx文件中的特定文本内容
"""

import sys
from pathlib import Path
from docx import Document
import logging
from typing import Dict

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def replace_text_in_paragraph(paragraph, replacements: Dict[str, str]) -> int:
    """在段落中替换文本"""
    replace_count = 0
    for old_text, new_text in replacements.items():
        if old_text in paragraph.text:
            # 处理段落中的runs
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    replace_count += 1
    return replace_count

def replace_text_in_table(table, replacements: Dict[str, str]) -> int:
    """在表格中替换文本"""
    replace_count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_count += replace_text_in_paragraph(paragraph, replacements)
    return replace_count

def process_single_document(file_path: str) -> bool:
    """处理单个文档"""
    replacement_map = {
        '生产部': '装配部',
        '市场部': '业务部', 
        '货仓': '仓库',
        '质安部': '品质部',
        '制造部': '五金、装配部',
        '车间组': '生产线',
        '班组': '拉线',
        '设计部': '研发部',
        '供销科': '业务部',
        '检验科': '品质部',
        '车间组长': '拉长',
        '修改员': '维修员',
        '调机员': '夹具设备维护员',
        '制模工': '夹具设备维护员',
        '领班': '拉长',
        '班组长': '拉长',
        '车间负责人': '生产主管',
        '仓务员': '仓管员',
        '付总': '副总'
    }
    
    try:
        logger.info(f"正在处理文件: {file_path}")
        doc = Document(file_path)
        total_replacements = 0
        
        # 处理段落
        for paragraph in doc.paragraphs:
            total_replacements += replace_text_in_paragraph(paragraph, replacement_map)
        
        # 处理表格
        for table in doc.tables:
            total_replacements += replace_text_in_table(table, replacement_map)
        
        # 保存文档
        doc.save(file_path)
        
        if total_replacements > 0:
            logger.info(f"文件处理完成，共替换 {total_replacements} 处")
        else:
            logger.info(f"文件无需替换")
            
        return True
        
    except Exception as e:
        logger.error(f"处理文件时出错: {str(e)}")
        return False

def main():
    """主函数"""
    file_path = "S:\\PG-GMO\\01-Input\\原始文档\\PG-ISO文件_docx\\HQ-QP-09 生产计划和生产过程控制程序\\HQ-QP-09 生产计划和生产过程控制程序.docx"
    
    success = process_single_document(file_path)
    
    if success:
        print("\n文件处理成功!")
    else:
        print("\n文件处理失败!")

if __name__ == "__main__":
    main()