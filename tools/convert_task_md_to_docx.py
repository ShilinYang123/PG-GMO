#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将工作任务清单中的Markdown文件转换为Word文档
作者：Qwen
"""

import os
import sys
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
import re

def setup_document_styles(doc):
    """设置文档样式"""
    # 设置标题样式
    try:
        heading1 = doc.styles['Heading 1']
        heading1.font.name = '微软雅黑'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        
        heading2 = doc.styles['Heading 2']
        heading2.font.name = '微软雅黑'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        
        heading3 = doc.styles['Heading 3']
        heading3.font.name = '微软雅黑'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        
        # 设置正文样式
        normal = doc.styles['Normal']
        normal.font.name = '宋体'
        normal.font.size = Pt(11)
    except Exception as e:
        print(f"设置样式时出错: {e}")

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """将单个markdown文件转换为docx文件"""
    try:
        # 读取markdown文件
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 创建Word文档
        doc = Document()
        setup_document_styles(doc)
        
        # 按行处理markdown内容
        lines = md_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                # 空行
                doc.add_paragraph()
                continue
            
            # 处理标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                p = doc.add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                p = doc.add_heading(line[7:], level=6)
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            # 处理表格（简单处理）
            elif '|' in line and line.count('|') >= 2:
                # 简单表格处理，这里只是添加为段落
                p = doc.add_paragraph(line)
            else:
                # 普通段落
                p = doc.add_paragraph(line)
        
        # 保存文档
        os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
        doc.save(docx_file_path)
        print(f"转换完成: {md_file_path} -> {docx_file_path}")
        return True
        
    except Exception as e:
        print(f"转换失败 {md_file_path}: {e}")
        return False

def convert_specific_files():
    """转换指定的文件"""
    # 指定的文件列表
    files_to_convert = [
        "工作任务05-LatenSync数字人落实探索.md",
        "工作任务06-EasySpider可视化爬虫工具评估.md",
        "工作任务07-Qwen-Image文生图模型评估.md",
        "工作任务08-MIRIX多智能体记忆系统评估.md",
        "工作任务09-FinGPT金融智能体评估.md",
        "工作任务10-WebDancer信息检索Agent评估.md",
        "工作任务11-React前端动画库评估.md",
        "工作任务12-VMACD量能指标评估.md",
        "工作任务13-微软Qlib AI量化投资平台评估.md",
        "工作任务14-全球首个AI设计Agent评估.md",
        "工作任务15-Reddit创意社区平台评估.md",
        "工作任务16-AI大模型生态系统MCP+Cline+A2A+n8n+KAG评估.md",
        "工作任务17-全球首个世界一致性模型AI视频评估.md"
    ]
    
    source_dir = "S:/PG-GMO/03-WorkTask/任务清单"
    target_dir = "S:/PG-GMO/02-Output"
    
    converted_count = 0
    failed_count = 0
    
    for file_name in files_to_convert:
        md_file_path = os.path.join(source_dir, file_name)
        # 将.md替换为.docx
        docx_file_name = file_name.replace('.md', '.docx')
        docx_file_path = os.path.join(target_dir, docx_file_name)
        
        # 转换文件
        if os.path.exists(md_file_path):
            if convert_markdown_to_docx(md_file_path, docx_file_path):
                converted_count += 1
            else:
                failed_count += 1
        else:
            print(f"文件不存在: {md_file_path}")
            failed_count += 1
    
    print(f"\n转换完成统计:")
    print(f"成功转换: {converted_count} 个文件")
    print(f"转换失败: {failed_count} 个文件")
    
    return converted_count, failed_count

def main():
    """主函数"""
    print("=== Markdown to Word 转换器 ===")
    print("开始转换工作任务清单中的指定文件...\n")
    
    # 开始转换
    converted, failed = convert_specific_files()
    
    if failed == 0:
        print("\n✅ 所有文件转换成功！")
    else:
        print(f"\n⚠️ 转换完成，但有 {failed} 个文件转换失败")
    
    print(f"转换后的文件保存在: S:/PG-GMO/02-Output")

if __name__ == "__main__":
    main()