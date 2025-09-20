#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word鏂囨。鍚堝苟鍣?- PG绠＄悊鎵嬪唽鐢熸垚鍣?浣滆€咃細闆ㄤ繆
鍔熻兘锛氬皢鎵€鏈夐儴闂ㄧ殑.docx鏂囦欢鍚堝苟涓轰竴涓畬鏁寸殑PG绠＄悊鎵嬪唽
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import re

def create_cover_page(doc):
    """鍒涘缓灏侀潰椤?""
    # 娣诲姞灏侀潰鏍囬
    title = doc.add_heading('', level=0)
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.text = 'PG绠＄悊鎵嬪唽'
    title_run.font.name = '寰蒋闆呴粦'
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞绌鸿
    for _ in range(8):
        doc.add_paragraph()
    
    # 娣诲姞鍓爣棰?    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉姹囩紪')
    subtitle_run.font.name = '寰蒋闆呴粦'
    subtitle_run.font.size = Pt(24)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞绌鸿
    for _ in range(10):
        doc.add_paragraph()
    
    # 娣诲姞鏃ユ湡
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'{datetime.now().strftime("%Y骞?m鏈?)}')
    date_run.font.name = '寰蒋闆呴粦'
    date_run.font.size = Pt(18)
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞鍒嗛〉绗?    doc.add_page_break()

def create_table_of_contents(doc, toc_entries):
    """鍒涘缓鐩綍椤?""
    # 鐩綍鏍囬
    toc_title = doc.add_heading('鐩綍', level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # 娣诲姞鐩綍鏉＄洰
    for entry in toc_entries:
        level = entry['level']
        title = entry['title']
        page = entry['page']
        
        toc_para = doc.add_paragraph()
        
        # 鏍规嵁绾у埆璁剧疆缂╄繘
        if level == 1:
            indent = 0
            font_size = Pt(14)
            is_bold = True
        elif level == 2:
            indent = 0.5
            font_size = Pt(12)
            is_bold = False
        else:
            indent = 1.0
            font_size = Pt(11)
            is_bold = False
        
        # 璁剧疆娈佃惤鏍煎紡
        toc_para.paragraph_format.left_indent = Inches(indent)
        
        # 娣诲姞鏍囬
        title_run = toc_para.add_run(title)
        title_run.font.name = '瀹嬩綋'
        title_run.font.size = font_size
        title_run.font.bold = is_bold
        
        # 娣诲姞鐐圭嚎
        dots_run = toc_para.add_run('.' * (60 - len(title) - len(str(page))))
        dots_run.font.name = '瀹嬩綋'
        dots_run.font.size = font_size
        
        # 娣诲姞椤电爜
        page_run = toc_para.add_run(str(page))
        page_run.font.name = '瀹嬩綋'
        page_run.font.size = font_size
    
    # 娣诲姞鍒嗛〉绗?    doc.add_page_break()

def setup_page_numbering(doc):
    """璁剧疆椤电爜"""
    try:
        # 鑾峰彇鎵€鏈夎妭
        sections = doc.sections
        
        for section in sections:
            # 璁剧疆椤佃剼
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 娣诲姞椤电爜瀛楁
            run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)
            
    except Exception as e:
        print(f"璁剧疆椤电爜鏃跺嚭閿? {e}")

def get_department_order():
    """瀹氫箟閮ㄩ棬椤哄簭"""
    return [
        '鎬荤粡鍔?,
        '琛屾斂閮?, 
        '璐㈠姟閮?,
        '鐮斿彂閮?,
        '涓氬姟閮?,
        '閲囪喘閮?,
        'PMC閮?,
        '浜旈噾閮?,
        '瑁呴厤閮?,
        '鍝佽川閮?
    ]

def merge_docx_files(source_dir, output_file):
    """鍚堝苟鎵€鏈塪ocx鏂囦欢"""
    source_path = Path(source_dir)
    
    if not source_path.exists():
        print(f"閿欒锛氭簮鐩綍涓嶅瓨鍦?{source_dir}")
        return False
    
    # 鍒涘缓涓绘枃妗?    main_doc = Document()
    
    # 璁剧疆鏂囨。鏍峰紡
    try:
        # 璁剧疆椤甸潰杈硅窛
        sections = main_doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    except Exception as e:
        print(f"璁剧疆椤甸潰杈硅窛鏃跺嚭閿? {e}")
    
    # 鍒涘缓灏侀潰
    create_cover_page(main_doc)
    
    # 鏀堕泦鐩綍淇℃伅
    toc_entries = []
    current_page = 3  # 灏侀潰鍜岀洰褰曞崰2椤碉紝鍐呭浠庣3椤靛紑濮?    
    # 鎸夐儴闂ㄩ『搴忓鐞嗘枃浠?    department_order = get_department_order()
    processed_departments = set()
    
    for dept_name in department_order:
        dept_dir = source_path / dept_name
        if dept_dir.exists() and dept_dir.is_dir():
            processed_departments.add(dept_name)
            
            # 娣诲姞閮ㄩ棬鏍囬
            dept_title = main_doc.add_heading(f'{dept_name}', level=1)
            dept_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 娣诲姞鍒扮洰褰?            toc_entries.append({
                'level': 1,
                'title': dept_name,
                'page': current_page
            })
            current_page += 1
            
            # 澶勭悊閮ㄩ棬涓嬬殑鎵€鏈塪ocx鏂囦欢
            docx_files = list(dept_dir.rglob('*.docx'))
            docx_files.sort(key=lambda x: x.name)  # 鎸夋枃浠跺悕鎺掑簭
            
            for docx_file in docx_files:
                try:
                    # 璇诲彇鏂囨。
                    sub_doc = Document(str(docx_file))
                    
                    # 鑾峰彇鏂囦欢鍚嶄綔涓虹珷鑺傛爣棰?                    file_title = docx_file.stem
                    
                    # 娣诲姞绔犺妭鏍囬
                    chapter_title = main_doc.add_heading(file_title, level=2)
                    
                    # 娣诲姞鍒扮洰褰?                    toc_entries.append({
                        'level': 2,
                        'title': file_title,
                        'page': current_page
                    })
                    
                    # 澶嶅埗鏂囨。鍐呭
                    for para in sub_doc.paragraphs:
                        if para.text.strip():  # 璺宠繃绌烘钀?                            new_para = main_doc.add_paragraph()
                            
                            # 澶嶅埗娈佃惤鏍煎紡鍜屽唴瀹?                            for run in para.runs:
                                new_run = new_para.add_run(run.text)
                                # 澶嶅埗瀛椾綋鏍煎紡
                                if run.font.name:
                                    new_run.font.name = run.font.name
                                if run.font.size:
                                    new_run.font.size = run.font.size
                                new_run.font.bold = run.font.bold
                                new_run.font.italic = run.font.italic
                    
                    # 娣诲姞琛ㄦ牸
                    for table in sub_doc.tables:
                        new_table = main_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = 'Table Grid'
                        
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_table.cell(i, j).text = cell.text
                    
                    current_page += max(1, len(sub_doc.paragraphs) // 20)  # 浼扮畻椤垫暟
                    
                    print(f"宸插悎骞? {docx_file.name}")
                    
                except Exception as e:
                    print(f"鍚堝苟鏂囦欢澶辫触 {docx_file}: {e}")
                    continue
            
            # 閮ㄩ棬涔嬮棿娣诲姞鍒嗛〉
            main_doc.add_page_break()
    
    # 澶勭悊鏈湪椤哄簭涓殑閮ㄩ棬
    for dept_dir in source_path.iterdir():
        if dept_dir.is_dir() and dept_dir.name not in processed_departments:
            dept_name = dept_dir.name
            
            # 娣诲姞閮ㄩ棬鏍囬
            dept_title = main_doc.add_heading(f'{dept_name}', level=1)
            dept_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 娣诲姞鍒扮洰褰?            toc_entries.append({
                'level': 1,
                'title': dept_name,
                'page': current_page
            })
            current_page += 1
            
            # 澶勭悊閮ㄩ棬涓嬬殑鎵€鏈塪ocx鏂囦欢
            docx_files = list(dept_dir.rglob('*.docx'))
            docx_files.sort(key=lambda x: x.name)
            
            for docx_file in docx_files:
                try:
                    sub_doc = Document(str(docx_file))
                    file_title = docx_file.stem
                    
                    chapter_title = main_doc.add_heading(file_title, level=2)
                    
                    toc_entries.append({
                        'level': 2,
                        'title': file_title,
                        'page': current_page
                    })
                    
                    for para in sub_doc.paragraphs:
                        if para.text.strip():
                            new_para = main_doc.add_paragraph()
                            for run in para.runs:
                                new_run = new_para.add_run(run.text)
                                if run.font.name:
                                    new_run.font.name = run.font.name
                                if run.font.size:
                                    new_run.font.size = run.font.size
                                new_run.font.bold = run.font.bold
                                new_run.font.italic = run.font.italic
                    
                    for table in sub_doc.tables:
                        new_table = main_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = 'Table Grid'
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_table.cell(i, j).text = cell.text
                    
                    current_page += max(1, len(sub_doc.paragraphs) // 20)
                    print(f"宸插悎骞? {docx_file.name}")
                    
                except Exception as e:
                    print(f"鍚堝苟鏂囦欢澶辫触 {docx_file}: {e}")
                    continue
            
            main_doc.add_page_break()
    
    # 鍦ㄥ皝闈㈠悗鎻掑叆鐩綍
    # 鐢变簬python-docx鐨勯檺鍒讹紝鎴戜滑闇€瑕侀噸鏂板垱寤烘枃妗ｆ潵姝ｇ‘鎻掑叆鐩綍
    final_doc = Document()
    
    # 璁剧疆椤甸潰鏍煎紡
    sections = final_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 鍒涘缓灏侀潰
    create_cover_page(final_doc)
    
    # 鍒涘缓鐩綍
    create_table_of_contents(final_doc, toc_entries)
    
    # 澶嶅埗涓昏鍐呭锛堣烦杩囧師鏉ョ殑灏侀潰锛?    skip_first_page = True
    for para in main_doc.paragraphs:
        if skip_first_page and ('PG绠＄悊鎵嬪唽' in para.text or '閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉姹囩紪' in para.text):
            continue
        if skip_first_page and para.text.strip() == '':
            continue
        if skip_first_page and any(dept in para.text for dept in department_order):
            skip_first_page = False
        
        if not skip_first_page:
            new_para = final_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
            
            # 澶嶅埗娈佃惤鏍峰紡
            if para.style.name.startswith('Heading'):
                new_para.style = para.style
            new_para.alignment = para.alignment
    
    # 澶嶅埗琛ㄦ牸
    for table in main_doc.tables:
        new_table = final_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = 'Table Grid'
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = cell.text
    
    # 璁剧疆椤电爜
    setup_page_numbering(final_doc)
    
    # 淇濆瓨鏂囨。
    try:
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        final_doc.save(output_file)
        print(f"\n鉁?绠＄悊鎵嬪唽鐢熸垚鎴愬姛: {output_file}")
        return True
    except Exception as e:
        print(f"淇濆瓨鏂囨。澶辫触: {e}")
        return False

def main():
    """涓诲嚱鏁?""
    print("=== PG绠＄悊鎵嬪唽鐢熸垚鍣?===")
    print("浣滆€咃細闆ㄤ繆")
    print("寮€濮嬪悎骞堕儴闂ㄦ枃妗?..\n")
    
    # 婧愮洰褰曞拰杈撳嚭鏂囦欢
    source_dir = "S:/PG-GMO/02-Output/閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉_docx"
    output_file = "S:/PG-GMO/02-Output/PG绠＄悊鎵嬪唽.docx"
    
    if not os.path.exists(source_dir):
        print(f"閿欒锛氭簮鐩綍涓嶅瓨鍦?{source_dir}")
        return
    
    # 寮€濮嬪悎骞?    success = merge_docx_files(source_dir, output_file)
    
    if success:
        print("\n馃帀 PG绠＄悊鎵嬪唽鐢熸垚瀹屾垚锛?)
        print(f"鏂囦欢浣嶇疆: {output_file}")
        print("\n鎵嬪唽鍖呭惈:")
        print("- 灏侀潰椤?)
        print("- 鐩綍椤?)
        print("- 鍚勯儴闂ㄥ畬鏁存枃妗?)
        print("- 椤电爜")
    else:
        print("\n鉂?绠＄悊鎵嬪唽鐢熸垚澶辫触")

if __name__ == "__main__":
    main()
