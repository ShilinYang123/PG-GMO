#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
鎵归噺鏀硅繘閮ㄩ棬绠＄悊鎵嬪唽鏍煎紡
浣滆€咃細闆ㄤ繆
鍔熻兘锛氭壒閲忎紭鍖栨墍鏈夐儴闂ㄧ鐞嗘墜鍐岀殑鏍煎紡鎺掔増
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import re
import glob

def setup_document_styles(doc):
    """璁剧疆鏂囨。鏍峰紡"""
    styles = doc.styles
    
    # 璁剧疆姝ｆ枃鏍峰紡
    try:
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '瀹嬩綋'
        normal_font.size = Pt(12)
        
        normal_para = normal_style.paragraph_format
        normal_para.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_para.line_spacing = 1.5
        normal_para.space_after = Pt(6)
        normal_para.first_line_indent = Pt(24)
    except:
        pass
    
    # 璁剧疆鏍囬鏍峰紡
    try:
        # 涓€绾ф爣棰?        heading1_style = styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = '寰蒋闆呴粦'
        heading1_font.size = Pt(18)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(0, 51, 102)
        
        heading1_para = heading1_style.paragraph_format
        heading1_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading1_para.space_before = Pt(18)
        heading1_para.space_after = Pt(12)
        heading1_para.keep_with_next = True
        
        # 浜岀骇鏍囬
        heading2_style = styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = '寰蒋闆呴粦'
        heading2_font.size = Pt(16)
        heading2_font.bold = True
        heading2_font.color.rgb = RGBColor(0, 51, 102)
        
        heading2_para = heading2_style.paragraph_format
        heading2_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading2_para.space_before = Pt(12)
        heading2_para.space_after = Pt(6)
        heading2_para.keep_with_next = True
        
        # 涓夌骇鏍囬
        heading3_style = styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = '寰蒋闆呴粦'
        heading3_font.size = Pt(14)
        heading3_font.bold = True
        heading3_font.color.rgb = RGBColor(51, 51, 51)
        
        heading3_para = heading3_style.paragraph_format
        heading3_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading3_para.space_before = Pt(12)
        heading3_para.space_after = Pt(6)
        heading3_para.first_line_indent = Pt(0)
        
    except Exception as e:
        print(f"璁剧疆鏍囬鏍峰紡鏃跺嚭閿? {e}")

def create_department_cover(doc, department_name):
    """鍒涘缓閮ㄩ棬涓撲笟灏侀潰"""
    section = doc.sections[0]
    section.top_margin = Inches(2)
    section.bottom_margin = Inches(2)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    
    # 涓绘爣棰?    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('PG绠＄悊鎵嬪唽')
    title_run.font.name = '寰蒋闆呴粦'
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 瑁呴グ绾?    line_para = doc.add_paragraph()
    line_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line_run = line_para.add_run('鈹? * 15)
    line_run.font.name = '寰蒋闆呴粦'
    line_run.font.size = Pt(12)
    line_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 绌鸿
    for _ in range(4):
        doc.add_paragraph()
    
    # 閮ㄩ棬鍚嶇О
    dept_para = doc.add_paragraph()
    dept_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_run = dept_para.add_run(department_name)
    dept_run.font.name = '寰蒋闆呴粦'
    dept_run.font.size = Pt(24)
    dept_run.font.bold = True
    dept_run.font.color.rgb = RGBColor(204, 0, 0)  # 绾㈣壊绐佸嚭閮ㄩ棬
    
    # 绌鸿
    for _ in range(6):
        doc.add_paragraph()
    
    # 鍓爣棰?    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉')
    subtitle_run.font.name = '寰蒋闆呴粦'
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # 绌鸿
    for _ in range(10):
        doc.add_paragraph()
    
    # 鏃ユ湡
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run(f'{datetime.now().strftime("%Y骞?m鏈?)} 鐗堟湰')
    date_run.font.name = '寰蒋闆呴粦'
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_page_break()

def create_department_toc(doc, toc_entries, department_name):
    """鍒涘缓閮ㄩ棬鐩綍"""
    # 鐩綍鏍囬
    toc_title = doc.add_heading(f'{department_name} - 鐩綍', level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # 娣诲姞鐩綍鏉＄洰
    for entry in toc_entries:
        level = entry['level']
        title = entry['title']
        page = entry['page']
        
        toc_para = doc.add_paragraph()
        
        if level == 1:
            indent = 0
            font_size = Pt(14)
            is_bold = True
            color = RGBColor(0, 51, 102)
        elif level == 2:
            indent = 0.5
            font_size = Pt(12)
            is_bold = False
            color = RGBColor(0, 0, 0)
        else:
            indent = 1.0
            font_size = Pt(11)
            is_bold = False
            color = RGBColor(102, 102, 102)
        
        toc_para.paragraph_format.left_indent = Inches(indent)
        toc_para.paragraph_format.space_after = Pt(3)
        
        # 鏍囬
        title_run = toc_para.add_run(title)
        title_run.font.name = '寰蒋闆呴粦'
        title_run.font.size = font_size
        title_run.font.bold = is_bold
        title_run.font.color.rgb = color
        
        # 鐐圭嚎
        dots_length = max(5, 45 - len(title) - len(str(page)))
        dots_run = toc_para.add_run('路' * dots_length)
        dots_run.font.name = '寰蒋闆呴粦'
        dots_run.font.size = Pt(10)
        dots_run.font.color.rgb = RGBColor(153, 153, 153)
        
        # 椤电爜
        page_run = toc_para.add_run(str(page))
        page_run.font.name = 'Times New Roman'
        page_run.font.size = font_size
        page_run.font.bold = is_bold
        page_run.font.color.rgb = color
    
    doc.add_page_break()

def setup_page_layout(doc):
    """璁剧疆椤甸潰甯冨眬"""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
        
        # 璁剧疆椤佃剼椤电爜
        footer = section.footer
        for para in footer.paragraphs:
            para.clear()
        
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

def format_paragraph_enhanced(para, new_para):
    """澧炲己娈佃惤鏍煎紡鍖?""
    if not para.text.strip():
        return
    
    text = para.text.strip()
    
    # 妫€鏌ュ垪琛ㄩ」
    if text.startswith(('鈥?, '路', '-', '*')) or re.match(r'^\d+[.)銆乚', text):
        new_para.paragraph_format.left_indent = Inches(0.5)
        new_para.paragraph_format.first_line_indent = Inches(-0.25)
        new_para.paragraph_format.space_after = Pt(3)
    else:
        new_para.paragraph_format.first_line_indent = Pt(24)
        new_para.paragraph_format.space_after = Pt(6)
    
    # 璁剧疆琛岃窛
    new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    new_para.paragraph_format.line_spacing = 1.5
    
    # 澶嶅埗鍐呭鍜屾牸寮?    for run in para.runs:
        new_run = new_para.add_run(run.text)
        new_run.font.name = run.font.name or '瀹嬩綋'
        new_run.font.size = run.font.size or Pt(12)
        new_run.font.bold = run.font.bold
        new_run.font.italic = run.font.italic
        new_run.font.underline = run.font.underline

def improve_department_manual(input_file, output_file):
    """鏀硅繘閮ㄩ棬绠＄悊鎵嬪唽鏍煎紡"""
    try:
        # 浠庢枃浠跺悕鎻愬彇閮ㄩ棬鍚嶇О
        filename = Path(input_file).stem
        department_name = filename.replace('PG绠＄悊鎵嬪唽-', '').replace('.docx', '')
        
        doc = Document(input_file)
        new_doc = Document()
        
        setup_document_styles(new_doc)
        setup_page_layout(new_doc)
        
        # 鏀堕泦鐩綍淇℃伅
        toc_entries = []
        current_page = 3
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                toc_entries.append({
                    'level': level,
                    'title': para.text,
                    'page': current_page
                })
                current_page += 1
        
        # 鍒涘缓灏侀潰鍜岀洰褰?        create_department_cover(new_doc, department_name)
        create_department_toc(new_doc, toc_entries, department_name)
        
        # 澶嶅埗鍐呭
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.split()[-1])
                    new_heading = new_doc.add_heading(para.text, level=level)
                else:
                    new_para = new_doc.add_paragraph()
                    format_paragraph_enhanced(para, new_para)
        
        # 澶嶅埗琛ㄦ牸
        for table in doc.tables:
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = 'Table Grid'
            
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    new_cell.text = cell.text
                    
                    for para in new_cell.paragraphs:
                        for run in para.runs:
                            run.font.name = '瀹嬩綋'
                            run.font.size = Pt(10)
        
        new_doc.save(output_file)
        print(f"鉁?{department_name} 鏍煎紡鏀硅繘瀹屾垚")
        return True
        
    except Exception as e:
        print(f"鉂?{input_file} 鏍煎紡鏀硅繘澶辫触: {e}")
        return False

def main():
    """涓诲嚱鏁?""
    print("=== 鎵归噺閮ㄩ棬绠＄悊鎵嬪唽鏍煎紡鏀硅繘鍣?===")
    print("浣滆€咃細闆ㄤ繆")
    print("寮€濮嬫壒閲忔敼杩涢儴闂ㄧ鐞嗘墜鍐屾牸寮?..\n")
    
    # 鏌ユ壘鎵€鏈夐儴闂ㄧ鐞嗘墜鍐?    input_dir = "S:/PG-GMO/02-Output"
    pattern = os.path.join(input_dir, "PG绠＄悊鎵嬪唽-*.docx")
    manual_files = glob.glob(pattern)
    
    if not manual_files:
        print("鏈壘鍒伴儴闂ㄧ鐞嗘墜鍐屾枃浠?)
        return
    
    success_count = 0
    total_count = len(manual_files)
    
    print(f"鎵惧埌 {total_count} 涓儴闂ㄧ鐞嗘墜鍐屾枃浠?)
    print("寮€濮嬫壒閲忓鐞?..\n")
    
    for input_file in manual_files:
        # 鐢熸垚杈撳嚭鏂囦欢鍚?        filename = Path(input_file).stem
        output_file = os.path.join(input_dir, f"{filename}_鏀硅繘鐗?docx")
        
        print(f"姝ｅ湪澶勭悊: {Path(input_file).name}")
        
        if improve_department_manual(input_file, output_file):
            success_count += 1
    
    print(f"\n=== 鎵归噺澶勭悊瀹屾垚 ===")
    print(f"鎬绘枃浠舵暟: {total_count}")
    print(f"鎴愬姛澶勭悊: {success_count}")
    print(f"澶辫触鏁伴噺: {total_count - success_count}")
    
    if success_count == total_count:
        print("\n馃帀 鎵€鏈夐儴闂ㄧ鐞嗘墜鍐屾牸寮忔敼杩涘畬鎴愶紒")
        print("\n鏀硅繘鍐呭:")
        print("- 涓撲笟閮ㄩ棬灏侀潰璁捐")
        print("- 涓€у寲鐩綍鏍煎紡")
        print("- 缁熶竴瀛椾綋鍜岄棿璺?)
        print("- 浼樺寲鏍囬灞傛")
        print("- 鏀硅繘娈佃惤鎺掔増")
        print("- 缇庡寲琛ㄦ牸鏍峰紡")
        print("- 瑙勮寖椤电爜鏄剧ず")
        print("\n鏀硅繘鐗堟枃浠跺凡淇濆瓨锛屾枃浠跺悕鍚庣紑'_鏀硅繘鐗?")
    else:
        print(f"\n鈿狅笍 閮ㄥ垎鏂囦欢澶勭悊澶辫触锛岃妫€鏌ラ敊璇俊鎭?)

if __name__ == "__main__":
    main()
