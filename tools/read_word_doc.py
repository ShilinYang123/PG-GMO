import docx
import sys

def read_word_document(file_path):
    try:
        # 打开Word文档
        doc = docx.Document(file_path)
        
        print(f"正在读取文档: {file_path}")
        print("=" * 50)
        
        # 读取所有段落
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 只打印非空段落
                print(f"段落 {i+1}: {paragraph.text}")
        
        print("\n" + "=" * 50)
        
        # 读取表格内容
        if doc.tables:
            print(f"文档包含 {len(doc.tables)} 个表格:")
            for table_idx, table in enumerate(doc.tables):
                print(f"\n表格 {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    print(f"  行 {row_idx + 1}: {' | '.join(row_data)}")
        
    except Exception as e:
        print(f"读取文档时出错: {e}")

if __name__ == "__main__":
    file_path = "S:\\PG-GMO\\office\\装配部\\生产部部门架构及工作职责.docx"
    read_word_document(file_path)