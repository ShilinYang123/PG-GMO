#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word文档转换器
作者：雨俊
功能：将部门基础建设工作成果中的所有.md文件转换为.docx文件
"""

import os
import sys
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
import re

def setup_document_styles(doc):
    """设置文档样式"""
    # 设置标题样式
    try:
        heading1 = doc.styles['Heading 1']
        heading1.font.name = '微软雅黑'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        
        heading2 = doc.styles['Heading 2']
        heading2.font.name = '微软雅黑'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        
        heading3 = doc.styles['Heading 3']
        heading3.font.name = '微软雅黑'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        
        # 设置正文样式
        normal = doc.styles['Normal']
        normal.font.name = '宋体'
        normal.font.size = Pt(11)
    except Exception as e:
        print(f"设置样式时出错: {e}")

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """将单个markdown文件转换为docx文件"""
    try:
        # 读取markdown文件
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 创建Word文档
        doc = Document()
        setup_document_styles(doc)
        
        # 按行处理markdown内容
        lines = md_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                # 空行
                doc.add_paragraph()
                continue
            
            # 处理标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                p = doc.add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                p = doc.add_heading(line[7:], level=6)
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            # 处理表格（简单处理）
            elif '|' in line and line.count('|') >= 2:
                # 简单表格处理，这里只是添加为段落
                p = doc.add_paragraph(line)
            else:
                # 普通段落
                p = doc.add_paragraph(line)
        
        # 保存文档
        os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
        doc.save(docx_file_path)
        print(f"转换完成: {md_file_path} -> {docx_file_path}")
        return True
        
    except Exception as e:
        print(f"转换失败 {md_file_path}: {e}")
        return False

def convert_all_md_files(source_dir, target_dir):
    """转换目录下所有的markdown文件"""
    source_path = Path(source_dir)
    target_path = Path(target_dir)
    
    converted_count = 0
    failed_count = 0
    
    # 遍历所有.md文件
    for md_file in source_path.rglob('*.md'):
        # 计算相对路径
        relative_path = md_file.relative_to(source_path)
        
        # 构建目标文件路径（将.md替换为.docx）
        docx_file = target_path / relative_path.with_suffix('.docx')
        
        # 转换文件
        if convert_markdown_to_docx(str(md_file), str(docx_file)):
            converted_count += 1
        else:
            failed_count += 1
    
    print(f"\n转换完成统计:")
    print(f"成功转换: {converted_count} 个文件")
    print(f"转换失败: {failed_count} 个文件")
    
    return converted_count, failed_count

def main():
    """主函数"""
    print("=== Markdown to Word 转换器 ===")
    print("作者：雨俊")
    print("开始转换部门基础建设工作成果文档...\n")
    
    # 源目录和目标目录
    source_dir = "S:/PG-GMO/Output/部门基础建设工作成果"
    target_dir = "S:/PG-GMO/Output/部门基础建设工作成果_docx"
    
    if not os.path.exists(source_dir):
        print(f"错误：源目录不存在 {source_dir}")
        return
    
    # 开始转换
    converted, failed = convert_all_md_files(source_dir, target_dir)
    
    if failed == 0:
        print("\n✅ 所有文件转换成功！")
    else:
        print(f"\n⚠️ 转换完成，但有 {failed} 个文件转换失败")
    
    print(f"转换后的文件保存在: {target_dir}")

if __name__ == "__main__":
    main()