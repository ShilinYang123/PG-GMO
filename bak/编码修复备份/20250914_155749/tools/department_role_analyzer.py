#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
品高ISO文档部门角色统计分析器
统计所有ISO文档中出现的部门和角色名称
"""

import os
import re
import json
from pathlib import Path
from collections import Counter, defaultdict
from datetime import datetime
import subprocess
import sys

class DepartmentRoleAnalyzer:
    def __init__(self, input_dir, output_dir):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 部门关键词模式
        self.department_patterns = [
            r'(\w*部)(?![门户])',  # 以"部"结尾的词
            r'(\w*室)(?![内外])',  # 以"室"结尾的词
            r'(\w*科)(?![学技])',  # 以"科"结尾的词
            r'(\w*组)(?![织成])',  # 以"组"结尾的词
            r'(\w*中心)',         # 以"中心"结尾的词
            r'(\w*办公室)',       # 办公室
            r'(总经办)',          # 总经办
            r'(管理层)',          # 管理层
            r'(高层)',            # 高层
        ]
        
        # 角色/职位关键词模式
        self.role_patterns = [
            r'(\w*经理)',         # 经理
            r'(\w*主管)',         # 主管
            r'(\w*负责人)',       # 负责人
            r'(\w*专员)',         # 专员
            r'(\w*员)',           # 员工
            r'(\w*师)',           # 工程师等
            r'(\w*长)',           # 部长、科长等
            r'(\w*总)',           # 总监、总经理等
            r'(\w*代表)',         # 代表
            r'(审核员)',          # 审核员
            r'(检验员)',          # 检验员
            r'(操作员)',          # 操作员
            r'(质检员)',          # 质检员
            r'(采购员)',          # 采购员
            r'(业务员)',          # 业务员
            r'(文控员)',          # 文控员
        ]
        
        # 需要排除的通用词汇
        self.exclude_words = {
            '部分', '部门', '全部', '内部', '外部', '局部', '头部', '尾部',
            '室内', '室外', '教室', '会议室', '办公室',
            '科学', '科技', '学科', '本科',
            '组织', '组成', '小组', '工作组',
            '中心思想', '市中心', '重心',
            '总结', '总计', '总和', '汇总',
            '长度', '长短', '成长', '增长',
            '工程师', '程序师', '设计师', '分析师'
        }
        
        self.departments = Counter()
        self.roles = Counter()
        self.document_stats = []
        
    def read_document_content(self, file_path):
        """读取文档内容"""
        try:
            if file_path.suffix.lower() == '.docx':
                # 直接使用docx库读取
                from docx import Document
                doc = Document(file_path)
                
                # 读取所有段落文本
                all_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        all_text.append(para.text.strip())
                
                # 读取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                all_text.append(cell.text.strip())
                
                return '\n'.join(all_text)
                
            elif file_path.suffix.lower() == '.doc':
                print(f"跳过.doc文件（不支持）: {file_path.name}")
                return ""
            elif file_path.suffix.lower() in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                print(f"不支持的文件格式: {file_path}")
                return ""
        except Exception as e:
            print(f"读取文档失败 {file_path}: {e}")
            return ""
    
    def extract_departments_and_roles(self, content, file_path):
        """从文档内容中提取部门和角色"""
        if not content:
            return [], []
        
        departments_found = set()
        roles_found = set()
        
        # 提取部门
        for pattern in self.department_patterns:
            matches = re.findall(pattern, content)
            for match in matches:
                if match and len(match) >= 2 and match not in self.exclude_words:
                    departments_found.add(match)
        
        # 提取角色
        for pattern in self.role_patterns:
            matches = re.findall(pattern, content)
            for match in matches:
                if match and len(match) >= 2 and match not in self.exclude_words:
                    roles_found.add(match)
        
        return list(departments_found), list(roles_found)
    
    def analyze_document(self, file_path):
        """分析单个文档"""
        print(f"正在分析: {file_path.name}")
        
        content = self.read_document_content(file_path)
        departments, roles = self.extract_departments_and_roles(content, file_path)
        
        # 更新计数器
        for dept in departments:
            self.departments[dept] += 1
        for role in roles:
            self.roles[role] += 1
        
        # 记录文档统计信息
        doc_stat = {
            'file_name': file_path.name,
            'file_path': str(file_path),
            'departments_count': len(departments),
            'roles_count': len(roles),
            'departments': departments,
            'roles': roles,
            'content_length': len(content) if content else 0
        }
        self.document_stats.append(doc_stat)
        
        return departments, roles
    
    def analyze_all_documents(self):
        """分析所有文档"""
        # 优先使用转换后的docx文件
        converted_dir = Path("S:/PG-GMO/01-Input/原始文档/PG-ISO文件_docx")
        
        print(f"开始分析目录: {self.input_dir}")
        print(f"转换文档目录: {converted_dir}")
        
        # 查找所有文档文件
        doc_files = []
        
        # 优先查找转换后的docx文件
        if converted_dir.exists():
            for root, dirs, files in os.walk(converted_dir):
                for file in files:
                    file_path = Path(root) / file
                    if file_path.suffix.lower() == '.docx':
                        doc_files.append(file_path)
            print(f"找到转换后的docx文件: {len(doc_files)} 个")
        
        # 如果没有转换文件，则使用原始文档文件
        if not doc_files:
            for root, dirs, files in os.walk(self.input_dir):
                for file in files:
                    file_path = Path(root) / file
                    if file_path.suffix.lower() in ['.doc', '.docx', '.txt', '.md']:
                        doc_files.append(file_path)
            print(f"找到原始文档文件: {len(doc_files)} 个")
        
        if not doc_files:
            print("未找到文档文件")
            print("建议先运行doc_to_docx_converter.py转换.doc文件")
            return self.generate_report()
        
        print(f"开始分析 {len(doc_files)} 个文档...")
        
        # 分析每个文档
        for file_path in doc_files:
            try:
                self.analyze_document(file_path)
            except Exception as e:
                print(f"分析文档失败 {file_path}: {e}")
        
        return self.generate_report()
    
    def generate_report(self):
        """生成分析报告"""
        report = {
            'analysis_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'total_documents': len(self.document_stats),
            'total_departments': len(self.departments),
            'total_roles': len(self.roles),
            'departments': dict(self.departments.most_common()),
            'roles': dict(self.roles.most_common()),
            'document_details': self.document_stats
        }
        
        # 保存JSON报告
        json_file = self.output_dir / 'ISO文档部门角色统计报告.json'
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        # 生成Markdown报告
        md_file = self.output_dir / 'ISO文档部门角色统计报告.md'
        self.generate_markdown_report(report, md_file)
        
        print(f"\n📊 分析完成!")
        print(f"📁 JSON报告: {json_file}")
        print(f"📄 Markdown报告: {md_file}")
        print(f"📈 统计结果:")
        print(f"   - 分析文档数: {report['total_documents']}")
        print(f"   - 发现部门数: {report['total_departments']}")
        print(f"   - 发现角色数: {report['total_roles']}")
        
        return report
    
    def generate_markdown_report(self, report, output_file):
        """生成Markdown格式报告"""
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("# 品高ISO文档部门角色统计分析报告\n\n")
            f.write(f"**分析时间**: {report['analysis_date']}\n")
            f.write(f"**分析文档数**: {report['total_documents']}\n")
            f.write(f"**发现部门数**: {report['total_departments']}\n")
            f.write(f"**发现角色数**: {report['total_roles']}\n\n")
            
            # 部门统计
            f.write("## 📊 部门统计\n\n")
            f.write("| 部门名称 | 出现次数 | 出现频率 |\n")
            f.write("|----------|----------|----------|\n")
            total_dept_mentions = sum(report['departments'].values())
            for dept, count in report['departments'].items():
                frequency = f"{count/total_dept_mentions*100:.1f}%"
                f.write(f"| {dept} | {count} | {frequency} |\n")
            
            # 角色统计
            f.write("\n## 👥 角色统计\n\n")
            f.write("| 角色名称 | 出现次数 | 出现频率 |\n")
            f.write("|----------|----------|----------|\n")
            total_role_mentions = sum(report['roles'].values())
            for role, count in report['roles'].items():
                frequency = f"{count/total_role_mentions*100:.1f}%"
                f.write(f"| {role} | {count} | {frequency} |\n")
            
            # 文档详情
            f.write("\n## 📋 文档分析详情\n\n")
            f.write("| 文档名称 | 部门数 | 角色数 | 内容长度 |\n")
            f.write("|----------|--------|--------|----------|\n")
            for doc in report['document_details']:
                f.write(f"| {doc['file_name']} | {doc['departments_count']} | {doc['roles_count']} | {doc['content_length']} |\n")
            
            # 部门详细分布
            f.write("\n## 🏢 部门详细分布\n\n")
            for doc in report['document_details']:
                if doc['departments']:
                    f.write(f"### {doc['file_name']}\n")
                    f.write("**部门**: " + ", ".join(doc['departments']) + "\n\n")
                    if doc['roles']:
                        f.write("**角色**: " + ", ".join(doc['roles']) + "\n\n")
            
            f.write("\n---\n")
            f.write("*报告由品高ISO文档部门角色统计分析器自动生成*\n")

def main():
    # 配置路径
    input_dir = "S:/PG-GMO/01-Input/原始文档/PG-ISO文件"
    output_dir = "S:/PG-GMO/02-Output"
    
    # 创建分析器并执行分析
    analyzer = DepartmentRoleAnalyzer(input_dir, output_dir)
    report = analyzer.analyze_all_documents()
    
    return report

if __name__ == "__main__":
    main()