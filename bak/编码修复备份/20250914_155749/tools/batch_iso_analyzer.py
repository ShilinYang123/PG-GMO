#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量ISO文档部门角色分析工具
用于重新分析所有ISO文档，确保统计的准确性和完整性

作者: 雨俊
创建时间: 2025-01-26
"""

import os
import json
import re
from pathlib import Path
from datetime import datetime
from collections import defaultdict, Counter
import sys

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

try:
    from tools.office_document_reader import read_office_document
except ImportError:
    print("警告: 无法导入office_document_reader，将使用基础文档读取功能")
    
    def read_office_document(file_path):
        """基础文档读取功能"""
        try:
            if file_path.endswith('.docx'):
                from docx import Document
                doc = Document(file_path)
                content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content.append(cell.text.strip())
                return '\n'.join(content)
            else:
                return f"不支持的文件格式: {file_path}"
        except Exception as e:
            return f"读取文件失败: {str(e)}"

class BatchISOAnalyzer:
    """批量ISO文档分析器"""
    
    def __init__(self):
        self.base_path = Path("S:/PG-GMO/01-Input/原始文档")
        self.output_path = Path("S:/PG-GMO/02-Output")
        
        # 扩展的部门关键词（基于HQ-QP-09的发现）
        self.department_keywords = [
            'PMC', 'QC', 'QA', '市场部', '工程部', '品质部', '货仓', '五金部', 
            '制造部', '生产部', '仓库', '行政部', '业务部', '采购部', '财务部',
            '总经办', '装配部', '供销科', '检验科', '文控中心', '设计部',
            '质安部', '项目部', '办公室', '试验室', '辨识部', '检查组',
            '内审组', '审核组', '内审小组', '品质部文控组', '车间组', '班组',
            '高层', '管理层', '内部', '外部', '内外部'
        ]
        
        # 扩展的角色关键词（基于HQ-QP-09的发现）
        self.role_keywords = [
            '总经理', '副总经理', '经理', '主管', '负责人', '技工', '操作员',
            '领班', '采购员', '物料员', '库管', '制模工', '检验员', '质检员',
            '仓管员', '仓务员', '跟单员', '专员', '拉长', '生产线拉长',
            '调机员', '工模师', '修理员', '生产员', '组长', '班组长',
            '车间负责人', '车间组长', '管理者代表', '技术副总', '付总',
            '副总', '人员', '员工', '工作人员', '技术人员', '管理人员',
            '操作人员', '检查员', '评价员', '评审员', '施工人员', '聘请老师'
        ]
        
        # 统计结果
        self.department_stats = Counter()
        self.role_stats = Counter()
        self.document_results = {}
        self.total_documents = 0
        self.processed_documents = 0
        self.failed_documents = []
        
    def find_all_documents(self):
        """查找所有ISO文档"""
        documents = []
        
        # 查找docx格式文档（优先）
        docx_path = self.base_path / "PG-ISO文件_docx"
        if docx_path.exists():
            for item in docx_path.rglob("*.docx"):
                if not item.name.startswith('~$'):  # 排除临时文件
                    documents.append(item)
        
        # 查找doc格式文档（作为备选）
        doc_path = self.base_path / "PG-ISO文件"
        if doc_path.exists():
            for item in doc_path.rglob("*.doc"):
                # 检查是否已有对应的docx文件
                docx_equivalent = docx_path / item.relative_to(doc_path).with_suffix('.docx')
                if not docx_equivalent.exists():
                    documents.append(item)
        
        return sorted(documents)
    
    def analyze_document_content(self, content, doc_name):
        """分析单个文档内容"""
        if not content or "读取文件失败" in content or "不支持的文件格式" in content:
            return {}, {}
        
        lines = content.split('\n')
        doc_departments = Counter()
        doc_roles = Counter()
        
        for line_num, line in enumerate(lines, 1):
            line = line.strip()
            if not line:
                continue
            
            # 分析部门 - 使用简单包含匹配（与专门分析工具一致）
            for dept in self.department_keywords:
                if dept in line:
                    # 计算出现次数
                    count = line.count(dept)
                    if count > 0:
                        doc_departments[dept] += count
                        print(f"  [{doc_name}] 第{line_num}行发现部门 '{dept}' {count}次: {line[:100]}...")
            
            # 分析角色 - 使用简单包含匹配（与专门分析工具一致）
            for role in self.role_keywords:
                if role in line:
                    # 计算出现次数
                    count = line.count(role)
                    if count > 0:
                        doc_roles[role] += count
                        print(f"  [{doc_name}] 第{line_num}行发现角色 '{role}' {count}次: {line[:100]}...")
        
        return doc_departments, doc_roles
    
    def analyze_single_document(self, doc_path):
        """分析单个文档"""
        doc_name = doc_path.name
        print(f"\n正在分析: {doc_name}")
        print(f"文件路径: {doc_path}")
        
        try:
            # 读取文档内容
            content = read_office_document(str(doc_path))
            
            if not content or "读取文件失败" in content:
                print(f"  ❌ 读取失败: {content}")
                self.failed_documents.append(str(doc_path))
                return
            
            print(f"  ✅ 成功读取，内容长度: {len(content)} 字符")
            
            # 分析内容
            doc_departments, doc_roles = self.analyze_document_content(content, doc_name)
            
            # 更新总统计
            self.department_stats.update(doc_departments)
            self.role_stats.update(doc_roles)
            
            # 保存文档结果
            self.document_results[doc_name] = {
                'path': str(doc_path),
                'departments': dict(doc_departments),
                'roles': dict(doc_roles),
                'content_length': len(content),
                'analysis_time': datetime.now().isoformat()
            }
            
            print(f"  📊 发现部门: {len(doc_departments)} 个")
            print(f"  👥 发现角色: {len(doc_roles)} 个")
            
            self.processed_documents += 1
            
        except Exception as e:
            print(f"  ❌ 分析失败: {str(e)}")
            self.failed_documents.append(str(doc_path))
    
    def run_analysis(self):
        """运行批量分析"""
        print("🚀 开始批量ISO文档分析...")
        print(f"基础路径: {self.base_path}")
        
        # 查找所有文档
        documents = self.find_all_documents()
        self.total_documents = len(documents)
        
        print(f"\n📁 找到 {self.total_documents} 个文档")
        
        if not documents:
            print("❌ 未找到任何文档！")
            return
        
        # 分析每个文档
        for doc_path in documents:
            self.analyze_single_document(doc_path)
        
        # 生成报告
        self.generate_reports()
        
        print(f"\n✅ 分析完成！")
        print(f"📊 总文档数: {self.total_documents}")
        print(f"✅ 成功处理: {self.processed_documents}")
        print(f"❌ 失败文档: {len(self.failed_documents)}")
        
        if self.failed_documents:
            print("\n失败的文档:")
            for failed in self.failed_documents:
                print(f"  - {failed}")
    
    def generate_reports(self):
        """生成分析报告"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 生成JSON报告
        json_report = {
            'analysis_info': {
                'timestamp': timestamp,
                'total_documents': self.total_documents,
                'processed_documents': self.processed_documents,
                'failed_documents': len(self.failed_documents),
                'analyzer_version': '2.0 - 批量重新分析版'
            },
            'department_statistics': dict(self.department_stats.most_common()),
            'role_statistics': dict(self.role_stats.most_common()),
            'document_details': self.document_results,
            'failed_documents': self.failed_documents
        }
        
        json_file = self.output_path / "ISO文档批量重新分析结果.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(json_report, f, ensure_ascii=False, indent=2)
        
        print(f"\n📄 JSON报告已保存: {json_file}")
        
        # 生成Markdown报告
        self.generate_markdown_report(timestamp)
    
    def generate_markdown_report(self, timestamp):
        """生成Markdown格式报告"""
        md_content = f"""# ISO文档批量重新分析报告 🔄

**分析时间**: {timestamp}
**分析工具**: 批量ISO分析器 v2.0
**总文档数**: {self.total_documents}
**成功处理**: {self.processed_documents}
**失败文档**: {len(self.failed_documents)}

## 🔥 重要说明

本次分析是对之前统计结果的全面重新检查，使用改进的分析工具确保准确性。
基于杨老师发现的HQ-QP-09统计遗漏问题，对所有文档进行了重新分析。

## 📊 部门统计总览

| 排名 | 部门名称 | 出现次数 | 占比 |
|------|----------|----------|------|
"""
        
        total_dept_count = sum(self.department_stats.values())
        for i, (dept, count) in enumerate(self.department_stats.most_common(20), 1):
            percentage = (count / total_dept_count * 100) if total_dept_count > 0 else 0
            md_content += f"| {i} | **{dept}** | {count} | {percentage:.1f}% |\n"
        
        md_content += f"\n**部门总计**: {len(self.department_stats)} 个不同部门\n"
        md_content += f"**部门提及总次数**: {total_dept_count} 次\n\n"
        
        md_content += "## 👥 角色统计总览\n\n"
        md_content += "| 排名 | 角色名称 | 出现次数 | 占比 |\n"
        md_content += "|------|----------|----------|------|\n"
        
        total_role_count = sum(self.role_stats.values())
        for i, (role, count) in enumerate(self.role_stats.most_common(20), 1):
            percentage = (count / total_role_count * 100) if total_role_count > 0 else 0
            md_content += f"| {i} | **{role}** | {count} | {percentage:.1f}% |\n"
        
        md_content += f"\n**角色总计**: {len(self.role_stats)} 个不同角色\n"
        md_content += f"**角色提及总次数**: {total_role_count} 次\n\n"
        
        # 添加文档详情
        md_content += "## 📋 文档分析详情\n\n"
        
        for doc_name, details in sorted(self.document_results.items()):
            md_content += f"### {doc_name}\n\n"
            
            if details['departments']:
                md_content += "**部门**: "
                dept_list = [f"{dept}({count}次)" for dept, count in sorted(details['departments'].items(), key=lambda x: x[1], reverse=True)]
                md_content += ", ".join(dept_list) + "\n\n"
            else:
                md_content += "**部门**: 未发现\n\n"
            
            if details['roles']:
                md_content += "**角色**: "
                role_list = [f"{role}({count}次)" for role, count in sorted(details['roles'].items(), key=lambda x: x[1], reverse=True)]
                md_content += ", ".join(role_list) + "\n\n"
            else:
                md_content += "**角色**: 未发现\n\n"
        
        if self.failed_documents:
            md_content += "## ❌ 处理失败的文档\n\n"
            for failed in self.failed_documents:
                md_content += f"- {failed}\n"
            md_content += "\n"
        
        md_content += f"## 📈 分析总结\n\n"
        md_content += f"- 本次重新分析发现了 **{len(self.department_stats)}** 个不同部门\n"
        md_content += f"- 本次重新分析发现了 **{len(self.role_stats)}** 个不同角色\n"
        md_content += f"- 部门提及总次数: **{total_dept_count}** 次\n"
        md_content += f"- 角色提及总次数: **{total_role_count}** 次\n"
        md_content += f"- 成功处理文档: **{self.processed_documents}/{self.total_documents}** 个\n\n"
        
        md_content += "---\n\n"
        md_content += "*本报告由批量ISO分析器自动生成*\n"
        
        md_file = self.output_path / "ISO文档批量重新分析报告.md"
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        print(f"📄 Markdown报告已保存: {md_file}")

def main():
    """主函数"""
    analyzer = BatchISOAnalyzer()
    analyzer.run_analysis()

if __name__ == "__main__":
    main()