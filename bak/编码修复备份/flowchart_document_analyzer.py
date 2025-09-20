#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
流程图文档分析器
用于分析ISO程序文件，识别适合生成流程图的文档

作者: 雨俊
创建时间: 2025年8月26日
"""

import os
import json
import logging
from pathlib import Path
from docx import Document
from datetime import datetime
import re

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('S:\\PG-GMO\\flowchart_analyzer.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class FlowchartDocumentAnalyzer:
    def __init__(self, source_dir):
        self.source_dir = Path(source_dir)
        self.analysis_results = []
        
        # 流程图适用性关键词
        self.flowchart_keywords = [
            '流程', '程序', '步骤', '操作', '审批', '审核', 
            '检验', '检查', '控制', '管理', '处理', '执行',
            '评审', '监控', '测量', '分析', '改进', '纠正',
            '识别', '评估', '应对', '变更', '召回', '投诉'
        ]
        
        # 不适合生成流程图的关键词
        self.exclude_keywords = [
            '表单', '表格', '记录表', '清单', '目录',
            '规范', '标准', '制度', '规定', '要求'
        ]
    
    def read_docx_content(self, file_path):
        """读取DOCX文档内容"""
        try:
            doc = Document(file_path)
            content = []
            
            # 读取段落内容
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 读取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text.strip())
            
            return '\n'.join(content)
        except Exception as e:
            logger.error(f"读取文档失败 {file_path}: {e}")
            return ""
    
    def analyze_document_suitability(self, file_path, content):
        """分析文档是否适合生成流程图"""
        filename = file_path.name
        
        # 基础评分
        score = 0
        reasons = []
        
        # 1. 文件名分析
        filename_lower = filename.lower()
        for keyword in self.flowchart_keywords:
            if keyword in filename_lower:
                score += 2
                reasons.append(f"文件名包含关键词: {keyword}")
        
        # 排除不适合的文档
        for exclude_word in self.exclude_keywords:
            if exclude_word in filename_lower:
                score -= 3
                reasons.append(f"文件名包含排除词: {exclude_word}")
        
        # 2. 内容分析
        content_lower = content.lower()
        
        # 检查流程相关内容
        flowchart_indicators = [
            r'第[一二三四五六七八九十\d]+步',
            r'步骤[\d一二三四五六七八九十]+',
            r'[\d一二三四五六七八九十]+\.[\d一二三四五六七八九十]+',
            r'流程图',
            r'操作流程',
            r'工作流程',
            r'审批流程',
            r'处理流程',
            r'控制流程'
        ]
        
        for pattern in flowchart_indicators:
            matches = re.findall(pattern, content)
            if matches:
                score += len(matches) * 1
                reasons.append(f"发现流程指示词: {pattern} ({len(matches)}次)")
        
        # 检查关键词频率
        for keyword in self.flowchart_keywords:
            count = content_lower.count(keyword)
            if count > 0:
                score += min(count, 5)  # 最多加5分
                if count >= 3:
                    reasons.append(f"关键词'{keyword}'出现{count}次")
        
        # 3. 文档结构分析
        lines = content.split('\n')
        numbered_lines = 0
        for line in lines:
            if re.match(r'^[\d一二三四五六七八九十]+[\.\.、]', line.strip()):
                numbered_lines += 1
        
        if numbered_lines >= 3:
            score += 3
            reasons.append(f"发现{numbered_lines}个编号条目")
        
        # 4. 特殊程序文件加分
        special_procedures = [
            '控制程序', '管理程序', '评审程序', '审核程序',
            '检验程序', '处理程序', '应对程序'
        ]
        
        for procedure in special_procedures:
            if procedure in filename:
                score += 3
                reasons.append(f"特殊程序文件: {procedure}")
        
        # 判断是否适合生成流程图
        suitable = score >= 5
        
        return {
            'suitable': suitable,
            'score': score,
            'reasons': reasons,
            'content_length': len(content),
            'line_count': len(lines)
        }
    
    def scan_documents(self):
        """扫描所有DOCX文档"""
        logger.info(f"开始扫描目录: {self.source_dir}")
        
        docx_files = list(self.source_dir.rglob('*.docx'))
        logger.info(f"找到 {len(docx_files)} 个DOCX文件")
        
        suitable_count = 0
        unsuitable_count = 0
        
        for file_path in docx_files:
            # 跳过临时文件
            if file_path.name.startswith('~$'):
                continue
            
            logger.info(f"分析文档: {file_path.name}")
            
            # 读取文档内容
            content = self.read_docx_content(file_path)
            
            if not content:
                logger.warning(f"文档内容为空: {file_path.name}")
                continue
            
            # 分析适用性
            analysis = self.analyze_document_suitability(file_path, content)
            
            # 记录分析结果
            result = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'relative_path': str(file_path.relative_to(self.source_dir)),
                'suitable_for_flowchart': analysis['suitable'],
                'suitability_score': analysis['score'],
                'reasons': analysis['reasons'],
                'content_length': analysis['content_length'],
                'line_count': analysis['line_count'],
                'analysis_time': datetime.now().isoformat()
            }
            
            self.analysis_results.append(result)
            
            if analysis['suitable']:
                suitable_count += 1
                logger.info(f"✅ 适合生成流程图: {file_path.name} (评分: {analysis['score']})")
            else:
                unsuitable_count += 1
                logger.info(f"❌ 不适合生成流程图: {file_path.name} (评分: {analysis['score']})")
        
        logger.info(f"扫描完成: 适合 {suitable_count} 个, 不适合 {unsuitable_count} 个")
        return self.analysis_results
    
    def generate_analysis_report(self, output_path):
        """生成分析报告"""
        suitable_docs = [doc for doc in self.analysis_results if doc['suitable_for_flowchart']]
        unsuitable_docs = [doc for doc in self.analysis_results if not doc['suitable_for_flowchart']]
        
        report = {
            'analysis_summary': {
                'total_documents': len(self.analysis_results),
                'suitable_documents': len(suitable_docs),
                'unsuitable_documents': len(unsuitable_docs),
                'analysis_time': datetime.now().isoformat()
            },
            'suitable_documents': suitable_docs,
            'unsuitable_documents': unsuitable_docs
        }
        
        # 保存JSON报告
        json_path = output_path.with_suffix('.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        # 生成Markdown报告
        self.generate_markdown_report(report, output_path)
        
        logger.info(f"分析报告已生成: {output_path}")
        return report
    
    def generate_markdown_report(self, report, output_path):
        """生成Markdown格式的分析报告"""
        md_content = f"""# 流程图文档分析报告

## 📊 分析概要

- **分析时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
- **文档总数**: {report['analysis_summary']['total_documents']} 个
- **适合生成流程图**: {report['analysis_summary']['suitable_documents']} 个
- **不适合生成流程图**: {report['analysis_summary']['unsuitable_documents']} 个
- **适用率**: {report['analysis_summary']['suitable_documents'] / report['analysis_summary']['total_documents'] * 100:.1f}%

## ✅ 适合生成流程图的文档

| 序号 | 文档名称 | 评分 | 主要原因 |
|------|----------|------|----------|
"""
        
        for i, doc in enumerate(report['suitable_documents'], 1):
            reasons = '; '.join(doc['reasons'][:3])  # 只显示前3个原因
            md_content += f"| {i} | {doc['file_name']} | {doc['suitability_score']} | {reasons} |\n"
        
        md_content += f"""

## ❌ 不适合生成流程图的文档

| 序号 | 文档名称 | 评分 | 主要原因 |
|------|----------|------|----------|
"""
        
        for i, doc in enumerate(report['unsuitable_documents'], 1):
            reasons = '; '.join(doc['reasons'][:3]) if doc['reasons'] else '评分过低'
            md_content += f"| {i} | {doc['file_name']} | {doc['suitability_score']} | {reasons} |\n"
        
        md_content += f"""

## 📋 详细分析结果

### 适合生成流程图的文档详情

"""
        
        for doc in report['suitable_documents']:
            md_content += f"""
#### {doc['file_name']}
- **路径**: {doc['relative_path']}
- **适用性评分**: {doc['suitability_score']}
- **内容长度**: {doc['content_length']} 字符
- **行数**: {doc['line_count']} 行
- **分析原因**:
"""
            for reason in doc['reasons']:
                md_content += f"  - {reason}\n"
            md_content += "\n"
        
        # 保存Markdown报告
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

def main():
    """主函数"""
    source_dir = "S:\\PG-GMO\\01-Input\\原始文档\\PG-ISO文件_docx"
    output_dir = Path("S:\\PG-GMO\\02-Output")
    output_dir.mkdir(exist_ok=True)
    
    # 创建分析器
    analyzer = FlowchartDocumentAnalyzer(source_dir)
    
    # 扫描文档
    results = analyzer.scan_documents()
    
    # 生成报告
    report_path = output_dir / "流程图文档分析报告.md"
    report = analyzer.generate_analysis_report(report_path)
    
    print(f"\n📊 分析完成!")
    print(f"总文档数: {report['analysis_summary']['total_documents']}")
    print(f"适合生成流程图: {report['analysis_summary']['suitable_documents']}")
    print(f"不适合生成流程图: {report['analysis_summary']['unsuitable_documents']}")
    print(f"报告已保存至: {report_path}")
    
    return report

if __name__ == "__main__":
    main()