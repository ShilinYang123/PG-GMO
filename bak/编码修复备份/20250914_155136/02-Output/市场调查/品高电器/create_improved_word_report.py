#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
鏀硅繘鐗堝競鍦鸿皟鏌ユ姤鍛奧ord鏍煎紡鐢熸垚鍣?
瑙ｅ喅鎺掔増闂锛屼娇鐢╓ord鍘熺敓鏍煎紡
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    """鍦ㄦ钀戒腑娣诲姞瓒呴摼鎺?""
    try:
        # 鑾峰彇娈佃惤鐨勭埗鏂囨。
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # 鍒涘缓瓒呴摼鎺ュ厓绱?
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # 鍒涘缓杩愯鍏冪礌
        new_run = OxmlElement('w:r')
        
        # 璁剧疆瓒呴摼鎺ユ牱寮?
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # 娣诲姞鏂囨湰
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
        return hyperlink
    except Exception as e:
        # 濡傛灉瓒呴摼鎺ュ垱寤哄け璐ワ紝娣诲姞鏅€氭枃鏈?
        paragraph.add_run(text)
        return None

def setup_document_styles(doc):
    """璁剧疆鏂囨。鏍峰紡"""
    # 璁剧疆姝ｆ枃鏍峰紡
    styles = doc.styles
    
    # 淇敼姝ｆ枃鏍峰紡
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '寰蒋闆呴粦'
    normal_font.size = Pt(11)
    
    # 璁剧疆鏍囬鏍峰紡
    for i in range(1, 5):
        heading_style = styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '寰蒋闆呴粦'
        heading_font.bold = True
        if i == 1:
            heading_font.size = Pt(18)
            heading_font.color.rgb = RGBColor(31, 73, 125)
        elif i == 2:
            heading_font.size = Pt(16)
            heading_font.color.rgb = RGBColor(31, 73, 125)
        elif i == 3:
            heading_font.size = Pt(14)
            heading_font.color.rgb = RGBColor(31, 73, 125)
        else:
            heading_font.size = Pt(12)
            heading_font.color.rgb = RGBColor(31, 73, 125)

def process_markdown_line(doc, line, list_level=0):
    """澶勭悊鍗曡Markdown鍐呭骞惰浆鎹负Word鏍煎紡"""
    line = line.strip()
    
    if not line:
        return
    
    # 澶勭悊鏍囬
    if line.startswith('# '):
        heading = doc.add_heading(line[2:], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif line.startswith('## '):
        heading = doc.add_heading(line[3:], level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif line.startswith('### '):
        heading = doc.add_heading(line[4:], level=3)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif line.startswith('#### '):
        heading = doc.add_heading(line[5:], level=4)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 澶勭悊鏃犲簭鍒楄〃
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        content = line[2:].strip()
        process_text_formatting(p, content)
    
    # 澶勭悊鏈夊簭鍒楄〃 - 浣跨敤Word鍘熺敓缂栧彿
    elif re.match(r'^\d+\. ', line):
        p = doc.add_paragraph(style='List Number')
        # 鎻愬彇鍒楄〃鍐呭锛堝幓鎺夋暟瀛楀拰鐐瑰彿锛?
        content = re.sub(r'^\d+\. ', '', line).strip()
        process_text_formatting(p, content)
    
    # 澶勭悊绮椾綋鏍囪鐨勭壒娈婃钀?
    elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
        p = doc.add_paragraph()
        run = p.add_run(line[2:-2])
        run.bold = True
        run.font.size = Pt(12)
    
    # 澶勭悊鏅€氭钀?
    else:
        # 璺宠繃Markdown鍏冩暟鎹拰鍒嗛殧绗?
        if (line.startswith('---') or 
            line.startswith('璋冩煡鏃ユ湡:') or 
            line.startswith('璋冩煡浜哄憳:') or
            line.startswith('**璋冩煡') or
            line == '---'):
            return
        
        p = doc.add_paragraph()
        process_text_formatting(p, line)

def process_text_formatting(paragraph, text):
    """澶勭悊鏂囨湰鏍煎紡鍖栵紙绮椾綋銆侀摼鎺ョ瓑锛?""
    # 澶勭悊瓒呴摼鎺?
    url_pattern = r'https?://[^\s]+'
    urls = re.findall(url_pattern, text)
    
    if urls:
        # 鍒嗗壊鏂囨湰鍜孶RL
        parts = re.split(url_pattern, text)
        url_index = 0
        
        for i, part in enumerate(parts):
            if part:
                # 澶勭悊绮椾綋鏂囨湰
                process_bold_text(paragraph, part)
            
            # 娣诲姞瓒呴摼鎺?
            if url_index < len(urls) and i < len(parts) - 1:
                add_hyperlink(paragraph, urls[url_index], urls[url_index])
                url_index += 1
    else:
        # 澶勭悊绮椾綋鏂囨湰
        process_bold_text(paragraph, text)

def process_bold_text(paragraph, text):
    """澶勭悊绮椾綋鏂囨湰鏍囪"""
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part:  # 璺宠繃绌哄瓧绗︿覆
                run = paragraph.add_run(part)
                if i % 2 == 1:  # 濂囨暟绱㈠紩鐨勯儴鍒嗘槸绮椾綋
                    run.bold = True
    else:
        paragraph.add_run(text)

def create_improved_word_report():
    """鍒涘缓鏀硅繘鐗圵ord鏍煎紡鐨勫競鍦鸿皟鏌ユ姤鍛?""
    
    # 璇诲彇Markdown鎶ュ憡鍐呭
    md_file_path = "S:\\PG-GMO\\\02-Output\\\甯傚満璋冩煡\\鍝侀珮鐢靛櫒\\宸ヤ綔浠诲姟-鍝侀珮2508.md"
    
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"閿欒锛氭壘涓嶅埌鏂囦欢 {md_file_path}")
        return
    
    # 鍒涘缓Word鏂囨。
    doc = Document()
    
    # 璁剧疆鏂囨。鏍峰紡
    setup_document_styles(doc)
    
    # 璁剧疆鏂囨。鏍囬
    title = doc.add_heading('姹熼棬甯傚搧楂樼數鍣ㄥ疄涓氭湁闄愬叕鍙稿皬瀹剁數甯傚満娣卞害璋冩煡鎶ュ憡', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 娣诲姞绌鸿
    doc.add_paragraph()
    
    # 瑙ｆ瀽Markdown鍐呭骞惰浆鎹负Word鏍煎紡
    lines = content.split('\n')
    in_code_block = False
    
    for line in lines:
        # 璺宠繃浠ｇ爜鍧?
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            continue
            
        # 璺宠繃绗竴琛屾爣棰橈紙宸茬粡鎵嬪姩娣诲姞锛?
        if line.strip().startswith('# 姹熼棬甯傚搧楂樼數鍣?):
            continue
            
        process_markdown_line(doc, line)
    
    # 璁剧疆椤甸潰杈硅窛
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 娣诲姞椤佃剼
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "姹熼棬甯傚搧楂樼數鍣ㄥ疄涓氭湁闄愬叕鍙?- 甯傚満璋冩煡鎶ュ憡 | HQ绾夸笂杩愯惀"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 淇濆瓨Word鏂囨。
    output_path = "S:\\PG-GMO\\\02-Output\\\甯傚満璋冩煡\\鍝侀珮鐢靛櫒\\鍝侀珮鐢靛櫒灏忓鐢靛競鍦鸿皟鏌ユ姤鍛奯鏀硅繘鐗?docx"
    doc.save(output_path)
    
    print(f"鉁?鏀硅繘鐗圵ord鏍煎紡鎶ュ憡宸茬敓鎴愶細{output_path}")
    print("馃搵 鏀硅繘鐗圭偣锛?)
    print("   - 浣跨敤Word鍘熺敓缂栧彿鍜岄」鐩鍙锋牸寮?)
    print("   - 涓撲笟鐨勬爣棰樻牱寮忓拰瀛椾綋璁剧疆")
    print("   - 鍚堥€傜殑娈佃惤闂磋窛鍜岄〉闈㈣竟璺?)
    print("   - 鏁版嵁鏉ユ簮閾炬帴鍙偣鍑昏烦杞?)
    print("   - 缁熶竴缂栧啓浜轰负HQ绾夸笂杩愯惀")
    print("   - 绗﹀悎鍟嗗姟鏂囨。鏍囧噯鐨勬帓鐗?)
    
    return output_path

if __name__ == "__main__":
    create_improved_word_report()
