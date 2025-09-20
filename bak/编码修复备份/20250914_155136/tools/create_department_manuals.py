#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
閮ㄩ棬绠＄悊鎵嬪唽鐢熸垚鍣?
浣滆€咃細闆ㄤ繆
鍔熻兘锛氫负姣忎釜閮ㄩ棬鍗曠嫭鐢熸垚涓€涓猈ord鏂囨。锛屾牸寮忎负'PG绠＄悊鎵嬪唽-XX閮?docx'
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import re

def create_cover_page(doc, department_name):
    """鍒涘缓閮ㄩ棬灏侀潰椤?""
    # 娣诲姞灏侀潰鏍囬
    title = doc.add_heading('', level=0)
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.text = f'PG绠＄悊鎵嬪唽'
    title_run.font.name = '寰蒋闆呴粦'
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞绌鸿
    for _ in range(4):
        doc.add_paragraph()
    
    # 娣诲姞閮ㄩ棬鍚嶇О
    dept_title = doc.add_heading('', level=0)
    dept_run = dept_title.runs[0] if dept_title.runs else dept_title.add_run()
    dept_run.text = f'{department_name}'
    dept_run.font.name = '寰蒋闆呴粦'
    dept_run.font.size = Pt(28)
    dept_run.font.bold = True
    dept_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞绌鸿
    for _ in range(6):
        doc.add_paragraph()
    
    # 娣诲姞鍓爣棰?
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉')
    subtitle_run.font.name = '寰蒋闆呴粦'
    subtitle_run.font.size = Pt(20)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞绌鸿
    for _ in range(8):
        doc.add_paragraph()
    
    # 娣诲姞鏃ユ湡
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'{datetime.now().strftime("%Y骞?m鏈?)}')
    date_run.font.name = '寰蒋闆呴粦'
    date_run.font.size = Pt(16)
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 娣诲姞鍒嗛〉绗?
    doc.add_page_break()

def create_department_toc(doc, file_list):
    """鍒涘缓閮ㄩ棬鐩綍椤?""
    # 鐩綍鏍囬
    toc_title = doc.add_heading('鐩綍', level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # 娣诲姞鐩綍鏉＄洰
    page_num = 3  # 浠庣3椤靛紑濮?
    for i, file_name in enumerate(file_list, 1):
        toc_para = doc.add_paragraph()
        
        # 璁剧疆娈佃惤鏍煎紡
        toc_para.paragraph_format.left_indent = Inches(0.5)
        
        # 娣诲姞搴忓彿鍜屾爣棰?
        title_text = f"{i}. {file_name}"
        title_run = toc_para.add_run(title_text)
        title_run.font.name = '瀹嬩綋'
        title_run.font.size = Pt(12)
        
        # 娣诲姞鐐圭嚎
        dots_count = max(1, 50 - len(title_text) - len(str(page_num)))
        dots_run = toc_para.add_run('.' * dots_count)
        dots_run.font.name = '瀹嬩綋'
        dots_run.font.size = Pt(12)
        
        # 娣诲姞椤电爜
        page_run = toc_para.add_run(str(page_num))
        page_run.font.name = '瀹嬩綋'
        page_run.font.size = Pt(12)
        
        page_num += 2  # 姣忎釜鏂囨。浼扮畻2椤?
    
    # 娣诲姞鍒嗛〉绗?
    doc.add_page_break()

def setup_page_numbering(doc):
    """璁剧疆椤电爜"""
    try:
        # 鑾峰彇鎵€鏈夎妭
        sections = doc.sections
        
        for section in sections:
            # 璁剧疆椤佃剼
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 娣诲姞椤电爜瀛楁
            run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)
            
    except Exception as e:
        print(f"璁剧疆椤电爜鏃跺嚭閿? {e}")

def create_department_manual(dept_dir, output_dir):
    """涓哄崟涓儴闂ㄥ垱寤虹鐞嗘墜鍐?""
    dept_name = dept_dir.name
    print(f"\n姝ｅ湪澶勭悊閮ㄩ棬: {dept_name}")
    
    # 鑾峰彇閮ㄩ棬涓嬬殑鎵€鏈塪ocx鏂囦欢
    docx_files = list(dept_dir.rglob('*.docx'))
    if not docx_files:
        print(f"  璀﹀憡: {dept_name} 娌℃湁鎵惧埌docx鏂囦欢")
        return False
    
    # 鎸夋枃浠跺悕鎺掑簭
    docx_files.sort(key=lambda x: x.name)
    
    # 鍒涘缓涓绘枃妗?
    main_doc = Document()
    
    # 璁剧疆鏂囨。鏍峰紡
    try:
        # 璁剧疆椤甸潰杈硅窛
        sections = main_doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    except Exception as e:
        print(f"  璁剧疆椤甸潰杈硅窛鏃跺嚭閿? {e}")
    
    # 鍒涘缓灏侀潰
    create_cover_page(main_doc, dept_name)
    
    # 鍒涘缓鐩綍
    file_names = [f.stem for f in docx_files]
    create_department_toc(main_doc, file_names)
    
    # 鍚堝苟閮ㄩ棬鏂囨。
    for docx_file in docx_files:
        try:
            # 璇诲彇鏂囨。
            sub_doc = Document(str(docx_file))
            
            # 鑾峰彇鏂囦欢鍚嶄綔涓虹珷鑺傛爣棰?
            file_title = docx_file.stem
            
            # 娣诲姞绔犺妭鏍囬
            chapter_title = main_doc.add_heading(file_title, level=1)
            chapter_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 澶嶅埗鏂囨。鍐呭
            for para in sub_doc.paragraphs:
                if para.text.strip():  # 璺宠繃绌烘钀?
                    new_para = main_doc.add_paragraph()
                    
                    # 澶嶅埗娈佃惤鏍煎紡鍜屽唴瀹?
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        # 澶嶅埗瀛椾綋鏍煎紡
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                    
                    # 澶嶅埗娈佃惤鏍峰紡
                    if para.style.name.startswith('Heading'):
                        new_para.style = para.style
                    new_para.alignment = para.alignment
            
            # 娣诲姞琛ㄦ牸
            for table in sub_doc.tables:
                new_table = main_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = 'Table Grid'
                
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
            
            # 鏂囨。涔嬮棿娣诲姞鍒嗛〉
            main_doc.add_page_break()
            
            print(f"  宸插悎骞? {docx_file.name}")
            
        except Exception as e:
            print(f"  鍚堝苟鏂囦欢澶辫触 {docx_file}: {e}")
            continue
    
    # 璁剧疆椤电爜
    setup_page_numbering(main_doc)
    
    # 淇濆瓨鏂囨。
    try:
        output_file = output_dir / f"PG绠＄悊鎵嬪唽-{dept_name}.docx"
        os.makedirs(output_dir, exist_ok=True)
        main_doc.save(str(output_file))
        print(f"  鉁?鐢熸垚鎴愬姛: {output_file}")
        return True
    except Exception as e:
        print(f"  鉂?淇濆瓨澶辫触: {e}")
        return False

def main():
    """涓诲嚱鏁?""
    print("=== 閮ㄩ棬绠＄悊鎵嬪唽鐢熸垚鍣?===")
    print("浣滆€咃細闆ㄤ繆")
    print("寮€濮嬩负姣忎釜閮ㄩ棬鐢熸垚鐙珛鐨勭鐞嗘墜鍐?..\n")
    
    # 婧愮洰褰曞拰杈撳嚭鐩綍
    source_dir = Path("S:/PG-GMO/02-Output/閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉_docx")
    output_dir = Path("S:/PG-GMO/02-Output")
    
    if not source_dir.exists():
        print(f"閿欒锛氭簮鐩綍涓嶅瓨鍦?{source_dir}")
        return
    
    # 缁熻淇℃伅
    success_count = 0
    total_count = 0
    
    # 閬嶅巻鎵€鏈夐儴闂ㄧ洰褰?
    for dept_dir in source_dir.iterdir():
        if dept_dir.is_dir():
            total_count += 1
            if create_department_manual(dept_dir, output_dir):
                success_count += 1
    
    # 杈撳嚭缁熻缁撴灉
    print(f"\n=== 鐢熸垚瀹屾垚缁熻 ===")
    print(f"鎬婚儴闂ㄦ暟: {total_count}")
    print(f"鎴愬姛鐢熸垚: {success_count}")
    print(f"澶辫触鏁伴噺: {total_count - success_count}")
    
    if success_count == total_count:
        print("\n馃帀 鎵€鏈夐儴闂ㄧ鐞嗘墜鍐岀敓鎴愭垚鍔燂紒")
    else:
        print(f"\n鈿狅笍 鏈?{total_count - success_count} 涓儴闂ㄧ敓鎴愬け璐?)
    
    print(f"\n鐢熸垚鐨勬枃浠朵繚瀛樺湪: {output_dir}")
    print("鏂囦欢鍛藉悕鏍煎紡: PG绠＄悊鎵嬪唽-XX閮?docx")

if __name__ == "__main__":
    main()
