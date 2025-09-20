#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word鏂囨。杞崲鍣?浣滆€咃細闆ㄤ繆
鍔熻兘锛氬皢閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉涓殑鎵€鏈?md鏂囦欢杞崲涓?docx鏂囦欢
"""

import os
import sys
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
import re

def setup_document_styles(doc):
    """璁剧疆鏂囨。鏍峰紡"""
    # 璁剧疆鏍囬鏍峰紡
    try:
        heading1 = doc.styles['Heading 1']
        heading1.font.name = '寰蒋闆呴粦'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        
        heading2 = doc.styles['Heading 2']
        heading2.font.name = '寰蒋闆呴粦'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        
        heading3 = doc.styles['Heading 3']
        heading3.font.name = '寰蒋闆呴粦'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        
        # 璁剧疆姝ｆ枃鏍峰紡
        normal = doc.styles['Normal']
        normal.font.name = '瀹嬩綋'
        normal.font.size = Pt(11)
    except Exception as e:
        print(f"璁剧疆鏍峰紡鏃跺嚭閿? {e}")

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """灏嗗崟涓猰arkdown鏂囦欢杞崲涓篸ocx鏂囦欢"""
    try:
        # 璇诲彇markdown鏂囦欢
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 鍒涘缓Word鏂囨。
        doc = Document()
        setup_document_styles(doc)
        
        # 鎸夎澶勭悊markdown鍐呭
        lines = md_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                # 绌鸿
                doc.add_paragraph()
                continue
            
            # 澶勭悊鏍囬
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                p = doc.add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                p = doc.add_heading(line[7:], level=6)
            # 澶勭悊鍒楄〃
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            # 澶勭悊琛ㄦ牸锛堢畝鍗曞鐞嗭級
            elif '|' in line and line.count('|') >= 2:
                # 绠€鍗曡〃鏍煎鐞嗭紝杩欓噷鍙槸娣诲姞涓烘钀?                p = doc.add_paragraph(line)
            else:
                # 鏅€氭钀?                p = doc.add_paragraph(line)
        
        # 淇濆瓨鏂囨。
        os.makedirs(os.path.dirname(docx_file_path), exist_ok=True)
        doc.save(docx_file_path)
        print(f"杞崲瀹屾垚: {md_file_path} -> {docx_file_path}")
        return True
        
    except Exception as e:
        print(f"杞崲澶辫触 {md_file_path}: {e}")
        return False

def convert_all_md_files(source_dir, target_dir):
    """杞崲鐩綍涓嬫墍鏈夌殑markdown鏂囦欢"""
    source_path = Path(source_dir)
    target_path = Path(target_dir)
    
    converted_count = 0
    failed_count = 0
    
    # 閬嶅巻鎵€鏈?md鏂囦欢
    for md_file in source_path.rglob('*.md'):
        # 璁＄畻鐩稿璺緞
        relative_path = md_file.relative_to(source_path)
        
        # 鏋勫缓鐩爣鏂囦欢璺緞锛堝皢.md鏇挎崲涓?docx锛?        docx_file = target_path / relative_path.with_suffix('.docx')
        
        # 杞崲鏂囦欢
        if convert_markdown_to_docx(str(md_file), str(docx_file)):
            converted_count += 1
        else:
            failed_count += 1
    
    print(f"\n杞崲瀹屾垚缁熻:")
    print(f"鎴愬姛杞崲: {converted_count} 涓枃浠?)
    print(f"杞崲澶辫触: {failed_count} 涓枃浠?)
    
    return converted_count, failed_count

def main():
    """涓诲嚱鏁?""
    print("=== Markdown to Word 杞崲鍣?===")
    print("浣滆€咃細闆ㄤ繆")
    print("寮€濮嬭浆鎹㈤儴闂ㄥ熀纭€寤鸿宸ヤ綔鎴愭灉鏂囨。...\n")
    
    # 婧愮洰褰曞拰鐩爣鐩綍
    source_dir = "S:/PG-GMO/02-Output/閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉"
    target_dir = "S:/PG-GMO/02-Output/閮ㄩ棬鍩虹寤鸿宸ヤ綔鎴愭灉_docx"
    
    if not os.path.exists(source_dir):
        print(f"閿欒锛氭簮鐩綍涓嶅瓨鍦?{source_dir}")
        return
    
    # 寮€濮嬭浆鎹?    converted, failed = convert_all_md_files(source_dir, target_dir)
    
    if failed == 0:
        print("\n鉁?鎵€鏈夋枃浠惰浆鎹㈡垚鍔燂紒")
    else:
        print(f"\n鈿狅笍 杞崲瀹屾垚锛屼絾鏈?{failed} 涓枃浠惰浆鎹㈠け璐?)
    
    print(f"杞崲鍚庣殑鏂囦欢淇濆瓨鍦? {target_dir}")

if __name__ == "__main__":
    main()
